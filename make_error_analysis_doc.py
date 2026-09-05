"""
สร้างเอกสาร Error Analysis -> docs/error_analysis.docx  (+ รูป figures/diagrams/ea_*.png)

    python make_error_analysis_doc.py

รวบรวม failure modes ที่พบจากการทดสอบระบบกับภาพจริง — ใช้ประกอบบท
"การอภิปรายผล / ข้อจำกัด" ของเล่ม เนื้อหาอิงหลักฐานที่วัดได้จริง (ระบุ path/ตัวเลข)

ต้องมี: python-docx, matplotlib, opencv-python  + โมเดล Stage 1/2 (สำหรับ regenerate รูป)
รูป ea_1 สร้างจาก pipeline จริง; ถ้าโหลดโมเดลไม่ได้จะข้ามรูปนั้น
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

BASE = Path(__file__).resolve().parent
DIA = BASE / "figures" / "diagrams"
DOCX = BASE / "docs" / "error_analysis.docx"

for _n in ("Tahoma", "Leelawadee UI", "Angsana New"):
    try:
        font_manager.findfont(_n, fallback_to_default=False)
        plt.rcParams["font.family"] = _n
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False

# ---- ตัวเลขหลักฐาน (วัดได้จริงในเซสชันทดสอบ 2026-09-05) ----
RUST_CONF = [
    ("NEU/Roboflow crop\n(สไตล์ชุดเทรน)", 0.92, "#2e7d32"),
    ("ภาพโซ่สนิม\nเต็มเฟรม", 0.21, "#e65100"),
    ("ภาพ texture สนิม\n(stock photo)", 0.02, "#b71c1c"),
]
RUST_THR = 0.90

STAGE1_RATIO = [
    ("NEU crop\ntest (n=416)", 0.05),
    ("ภาพฉาก\nประตูเหล็ก", 0.24),
    ("แผ่นเหล็กทาสี\nลานขยะ", 0.03),
    ("แผ่นสแตนเลส\n+ กราฟิกโฆษณา", 0.00),
]
STAGE1_FALLBACK_THR = 0.05


def make_figs():
    DIA.mkdir(parents=True, exist_ok=True)

    # ea_2 — rust confidence เทียบเกณฑ์
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    labels = [x[0] for x in RUST_CONF]
    vals = [x[1] for x in RUST_CONF]
    cols = [x[2] for x in RUST_CONF]
    b = ax.bar(labels, vals, color=cols, width=0.55)
    ax.axhline(RUST_THR, ls="--", color="#333", lw=1.2)
    ax.text(2.35, RUST_THR + 0.02, f"per-class threshold ของ rust = {RUST_THR:.2f}",
            ha="right", fontsize=8.5, color="#333")
    for r, v in zip(b, vals):
        ax.text(r.get_x() + r.get_width() / 2, v + 0.02, f"{v:.2f}", ha="center", fontsize=9)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("confidence ดิบสูงสุด (คลาส rust)")
    ax.set_title("Stage 2 — ความมั่นใจต่อภาพสนิมต่างสไตล์ (โมเดลเดียวกัน)")
    fig.tight_layout()
    fig.savefig(DIA / "ea_2_rust_conf.png", dpi=165)
    plt.close(fig)
    print("  ", (DIA / "ea_2_rust_conf.png").relative_to(BASE))

    # ea_3 — Stage 1 metal coverage
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    labels = [x[0] for x in STAGE1_RATIO]
    vals = [x[1] for x in STAGE1_RATIO]
    b = ax.bar(labels, vals, color="#1565c0", width=0.55)
    ax.axhline(STAGE1_FALLBACK_THR, ls="--", color="#c62828", lw=1.2)
    ax.text(3.4, STAGE1_FALLBACK_THR + 0.006,
            f"< {STAGE1_FALLBACK_THR:.2f} → fallback ตรวจทั้งภาพ", ha="right",
            fontsize=8.5, color="#c62828")
    for r, v in zip(b, vals):
        ax.text(r.get_x() + r.get_width() / 2, v + 0.006, f"{v:.0%}", ha="center", fontsize=9)
    ax.set_ylim(0, 0.30)
    ax.set_ylabel("สัดส่วนพื้นที่ที่ Stage 1 ระบุว่าเป็นโลหะ")
    ax.set_title("Stage 1 (DMS46) — พื้นที่โลหะที่ตรวจพบ ต่อชนิดภาพ")
    fig.tight_layout()
    fig.savefig(DIA / "ea_3_stage1_ratio.png", dpi=165)
    plt.close(fig)
    print("  ", (DIA / "ea_3_stage1_ratio.png").relative_to(BASE))

    # ea_1 — logo false positive (รันจริงถ้าทำได้)
    try:
        import cv2
        import app as A
        p = BASE / "real_test" / "images" / "real_002_stainless_scratch.webp"
        if p.exists():
            img = cv2.imread(str(p))
            rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            annotated, _s1, _v, rows = A.analyze(rgb, 0.15, False)
            cv2.imwrite(str(DIA / "ea_1_logo_fp.png"), cv2.cvtColor(annotated, cv2.COLOR_RGB2BGR))
            print("  ", (DIA / "ea_1_logo_fp.png").relative_to(BASE), "  rows:", rows)
    except Exception as e:
        print("   (ข้าม ea_1 — โหลดโมเดลไม่ได้:", e, ")")


MODES = [
    ("EA-1", "Domain gap ของ Stage 2 (สไตล์ภาพนอกการกระจายของชุดเทรน)",
     "Stage 2 เทรนบน NEU-DET (crop texture เกรย์สเกลระยะใกล้) + Roboflow rust/crack "
     "(ภาพสไตล์เฉพาะ) พอเจอภาพสนิมสไตล์อื่น—เช่น texture stock photo หรือโซ่สนิมเต็มเฟรม—"
     "ความมั่นใจตกจาก ~0.92 (สไตล์ชุดเทรน) เหลือ 0.02–0.21 หลุดใต้ threshold ทุกกรณี "
     "จึงรายงานว่า \"ไม่พบตำหนิ\" ทั้งที่ภาพเป็นสนิมชัดเจน",
     "ea_2_rust_conf.png",
     "หลักฐาน: ทดสอบ 3 ภาพ (สนิม.jpg raw conf ≈ 0.02 ; images(2).jpg โซ่สนิม 0.21 ; "
     "Danger-Rust crop 0.92) — โมเดลตัวเดียวกัน ต่างกันแค่สไตล์ภาพ"),
    ("EA-2", "Stage 1 (DMS46) หาโลหะไม่เจอบนเหล็กทาสี / สนิมหนัก / ฉากรก",
     "DMS46 เป็น material segmentation ระดับฉากทั่วไป ไม่ได้ fine-tune กับเหล็กอุตสาหกรรม "
     "บน test set (crop NEU 416 ภาพ) fallback_rate 78%, gt_area_kept 16%; บนภาพฉากจริง "
     "แผ่นเหล็กทาสีส้มในลานขยะได้ metal_ratio แค่ 3% (→ fallback ทั้งภาพ) "
     "เท่ากับ Stage 1 ไม่ได้ช่วยกรองพื้นหลังในเคสที่ควรช่วยที่สุด",
     "ea_3_stage1_ratio.png",
     "หลักฐาน: results/stage1_dms46_test.json (aggregate) + ทดสอบภาพฉาก 3 แบบในเซสชัน"),
    ("EA-3", "False positive จากกราฟิก/โลโก้/ลายน้ำในภาพจริง",
     "ภาพจากผู้ขายเหล็ก (โฆษณา) มักมีโลโก้บริษัท ตัวหนังสือ กรอบสี ทับบนภาพ "
     "Stage 1 หาโลหะไม่เจอ (metal_ratio 0%) → fallback ทั้งภาพ → Stage 2 ไปตีกรอบ "
     "\"โลโก้วงกลม\" แล้วจำแนกเป็น crack 37.7% (ผ่าน threshold ของ crack ที่ 0.36) "
     "ขณะที่รอยขีดข่วนจริงกลางภาพได้แค่ 0.17 (ไม่ผ่าน) → ระบบตอบผิดทั้งชนิดและตำแหน่ง",
     "ea_1_logo_fp.png",
     "หลักฐาน: real_test/images/real_002_stainless_scratch.webp — กรอบแดงอยู่ที่โลโก้ S.T.K. METAL"),
    ("EA-4", "Train/test leakage ในคลาส rust (แก้แล้ว — ดู check_leakage.py)",
     "ชุด Roboflow \"Danger-Rust\" เป็นภาพถ่ายรัว/เฟรมติดกัน การ split แบบสุ่มเดิมทำให้ "
     "rust ใน valid 100% / test 98% มีภาพเกือบเหมือนอยู่ใน train → rust mAP 0.995 และ "
     "overall mAP50 0.853 สูงเกินจริง แก้ด้วย group-aware re-split (resplit_grouped.py) "
     "→ leakage = 0 แล้ว retrain/วัดผลใหม่",
     None,
     "หลักฐาน: results/leakage_gray.json (932 คู่ก่อนแก้) / leakage_gray_after.json (0)"),
    ("EA-5", "คลาสที่อ่อนแม้บน benchmark: crack และ inclusion",
     "แม้บนชุดทดสอบสะอาด crack ได้ mAP50 ~0.67 / recall ~0.66 และ inclusion recall ~0.66 "
     "— crack: grayscale ลด contrast ของรอยแยก + geometry เส้นบางทำ IoU ต่ำ; "
     "inclusion: จุดเล็กกระจาย annotate ยาก เป็นข้อจำกัดที่พบใน literature ของ NEU-DET เช่นกัน "
     "(ตัวเลขรายคลาสสุดท้ายให้ยึดผลหลัง retrain บน split ใหม่)",
     None,
     "หลักฐาน: results/stage2_train-gray-s.json (per_class) — จะอัปเดตหลัง retrain"),
]

MITIGATION = [
    ("EA-1", "เพิ่ม augmentation ตอนเทรน (blur, JPEG artifact, brightness/contrast, "
             "overlay สังเคราะห์) ; ประเมิน cross-dataset (GC10-DET / Severstal) ; "
             "ระบุขอบเขตว่าใช้กับภาพสไตล์คัดกรอง (มือถือ ระยะ ~0.3–1 m) เท่านั้น"),
    ("EA-2", "แทน DMS46 ด้วย classifier/detector เหล็ก-ไม่เหล็กตัวเล็กที่เทรนกับโดเมนนี้ ; "
             "หรือคง soft-gate + fallback ปัจจุบันและรายงาน Stage 1 เป็น negative ablation"),
    ("EA-3", "crop ขอบภาพ/ตัดแถบข้อความอัตโนมัติก่อนเข้าระบบ ; เพิ่ม negative sample "
             "(โลโก้/ข้อความ/ฉาก) ตอนเทรน Stage 2 ; ยกเกณฑ์ความมั่นใจขั้นต่ำรวมเมื่อ fallback"),
    ("EA-4", "แก้แล้ว — group-aware re-split ; ผนวก check_leakage.py เข้าขั้นตอนเตรียมข้อมูล"),
    ("EA-5", "เก็บ crack แบบ RGB (ไม่ทำ grayscale) เป็น ablation ; oversample inclusion เพิ่ม ; "
             "ทดลอง imgsz สูงขึ้นเฉพาะ 2 คลาสนี้"),
]


def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Tahoma"
    doc.styles["Normal"].font.size = Pt(11)

    def h1(t):
        doc.add_page_break()
        doc.add_heading(t, level=1)

    def img(name, width=6.0, caption=None):
        p = DIA / name
        if not p.exists():
            doc.add_paragraph(f"[ไม่พบรูป {name} — รัน make_error_analysis_doc.py ใหม่หลังโมเดลพร้อม]")
            return
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Inches(width))
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size = Pt(9)
            c.runs[0].font.italic = True

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except KeyError:
            t.style = "Table Grid"
        for i, hh in enumerate(headers):
            t.rows[0].cells[i].text = hh
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)

    ti = doc.add_heading("การวิเคราะห์ข้อผิดพลาดของระบบ (Error Analysis)", 0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage: DMS46 → YOLO11s)"
                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารนี้รวบรวมรูปแบบความล้มเหลว (failure mode) ที่พบจากการทดสอบระบบกับภาพนอกชุด "
        "benchmark — ทั้งภาพถ่ายฉากจริงและภาพจากผู้ขายเหล็ก — เพื่อใช้ประกอบบท "
        "\"การอภิปรายผลและข้อจำกัด\" ทุกข้ออ้างอิงหลักฐานที่วัดได้จริง (ระบุไฟล์/ตัวเลข). "
        "ปรับปรุงล่าสุด 2026-09-05.")

    h1("1. สรุปรูปแบบความล้มเหลว")
    table(["รหัส", "รูปแบบ", "ผลกระทบ"], [
        ("EA-1", "Domain gap ของ Stage 2", "พลาดตำหนิจริงบนภาพสไตล์นอกชุดเทรน (recall ตก)"),
        ("EA-2", "Stage 1 หาโลหะไม่เจอ (เหล็กทาสี/ฉากรก)", "Stage 1 ไม่ช่วยกรองพื้นหลังในเคสที่ควรช่วย"),
        ("EA-3", "False positive จากโลโก้/ข้อความ", "ตอบผิดชนิดและตำแหน่ง"),
        ("EA-4", "Train/test leakage (rust) — แก้แล้ว", "metric เดิมสูงเกินจริง"),
        ("EA-5", "คลาสอ่อน: crack, inclusion", "recall ต่ำแม้บน benchmark"),
    ], widths=[0.6, 2.6, 3.0])

    for code, title, body, fig, evid in MODES:
        h1(f"{code}. {title}")
        doc.add_paragraph(body)
        if fig:
            img(fig, 6.0)
        p = doc.add_paragraph()
        r = p.add_run(evid)
        r.font.size = Pt(9)
        r.font.italic = True

    h1("7. แนวทางบรรเทา (Mitigation)")
    table(["รหัส", "แนวทาง"], MITIGATION, widths=[0.6, 5.8])
    doc.add_paragraph()
    doc.add_paragraph(
        "หมายเหตุ: EA-1/EA-2/EA-3 เป็นข้อจำกัดเชิงโดเมน ไม่ขึ้นกับปัญหา split — ยังใช้ได้แม้หลัง "
        "retrain ; EA-4 แก้แล้ว ; EA-5 ให้ยึดตัวเลขรายคลาสหลัง retrain บน split ใหม่")

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX))
    print("\nเขียน:", DOCX.relative_to(BASE))


if __name__ == "__main__":
    print("สร้างรูป Error Analysis...")
    make_figs()
    build_docx()
