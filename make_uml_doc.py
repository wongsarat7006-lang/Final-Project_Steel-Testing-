"""
สร้างเอกสาร UML / SA รวม 10 ไดอาแกรม -> docs/uml_sa_diagrams.docx
+ รูป PNG ใน figures/diagrams/uml_*.png

    python make_uml_doc.py

ต้องมี: python-docx, matplotlib

เนื้อหาไดอาแกรมอิงโค้ดจริง (pipeline.py / app.py / evaluate*.py / schema ของ *_result.json).
ระบบเป็น CLI + prototype UI (Gradio) + โมเดล ML — ไม่มี DB / login / บัญชีผู้ใช้:
  - ERD  = Conceptual Data Model (persist เป็นไฟล์ JSON/CSV/label ไม่ใช่ตาราง DB)
  - Class = โค้ด functional -> โมดูลใช้ stereotype «utility» (คลาสที่มีแต่ static operation)
  - State = วงจรชีวิตของ "ภาพ 1 รูป" ระหว่างถูกประมวลผลใน memory
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import (FancyBboxPatch, FancyArrowPatch, Ellipse, Circle,
                                Rectangle, Polygon)
from matplotlib.lines import Line2D
from matplotlib import font_manager

BASE = Path(__file__).resolve().parent
DIA = BASE / "figures" / "diagrams"
DOCX = BASE / "docs" / "uml_sa_diagrams.docx"

for _n in ("Tahoma", "Leelawadee UI", "Angsana New"):
    try:
        font_manager.findfont(_n, fallback_to_default=False)
        plt.rcParams["font.family"] = _n
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False

CL = {"a": "#E8EAF0", "s1": "#CFE2F3", "s2": "#D9EAD3", "data": "#FCE5CD",
      "out": "#EAD1DC", "note": "#FFF2CC", "ext": "#F0F0F0", "sys": "#DAE8FC"}


def new(Y, w=11):
    fig, ax = plt.subplots(figsize=(w, max(3.5, w * Y / 100 * 0.92)))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, Y)
    ax.axis("off")
    return fig, ax


def title(ax, t, sub=None):
    Y = ax.get_ylim()[1]
    ax.text(50, Y - 1.5, t, ha="center", va="top", fontsize=12, fontweight="bold")
    if sub:
        ax.text(50, Y - 5.6, sub, ha="center", va="top", fontsize=7.4, color="#666")


def box(ax, x, y, w, h, text, fc="#FFFFFF", fs=8.5, round_=True, ec="#555"):
    style = "round,pad=0.6,rounding_size=1.6" if round_ else "square,pad=0"
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle=style, fc=fc, ec=ec, lw=1.1))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, zorder=5)
    return {"cx": x + w / 2, "b": y, "t": y + h, "l": x, "r": x + w,
            "my": y + h / 2, "w": w, "h": h}


def oval(ax, cx, cy, w, h, text, fc=CL["s2"], fs=8):
    ax.add_patch(Ellipse((cx, cy), w, h, fc=fc, ec="#555", lw=1.1))
    ax.text(cx, cy, text, ha="center", va="center", fontsize=fs, zorder=5)
    return {"cx": cx, "b": cy - h / 2, "t": cy + h / 2, "l": cx - w / 2,
            "r": cx + w / 2, "my": cy}


def actor(ax, x, y, name, s=5.5):
    ax.add_patch(Circle((x, y + s * 1.7), s * 0.42, fc="white", ec="#333", lw=1.3, zorder=4))
    ax.plot([x, x], [y + s * 1.28, y + s * 0.35], "k-", lw=1.3)
    ax.plot([x - s * 0.7, x + s * 0.7], [y + s * 0.95, y + s * 0.95], "k-", lw=1.3)
    ax.plot([x, x - s * 0.55], [y + s * 0.35, y - s * 0.45], "k-", lw=1.3)
    ax.plot([x, x + s * 0.55], [y + s * 0.35, y - s * 0.45], "k-", lw=1.3)
    ax.text(x, y - s * 1.0, name, ha="center", va="center", fontsize=7.6, fontweight="bold")
    return {"cx": x, "my": y + s * 0.7, "l": x, "r": x}


def diamond(ax, cx, cy, w, h, text, fc=CL["note"], fs=7.0):
    ax.add_patch(Polygon([(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2),
                          (cx - w / 2, cy)], closed=True, fc=fc, ec="#555", lw=1.1))
    ax.text(cx, cy, text, ha="center", va="center", fontsize=fs, zorder=5)
    return {"cx": cx, "b": cy - h / 2, "t": cy + h / 2, "l": cx - w / 2,
            "r": cx + w / 2, "my": cy}


def arrow(ax, p1, p2, dashed=False, txt=None, fs=6.8, tx_dy=1.8, color="#333",
          head="-|>"):
    ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle=head, mutation_scale=12, lw=1.1,
                                 color=color, ls="--" if dashed else "-",
                                 shrinkA=0, shrinkB=2))
    if txt:
        mx, my = (p1[0] + p2[0]) / 2, (p1[1] + p2[1]) / 2 + tx_dy
        ax.text(mx, my, txt, fontsize=fs, ha="center", va="center", color="#333",
                bbox=dict(fc="white", ec="none", pad=0.5), zorder=6)


def dep(ax, p1, p2, txt=None, fs=6.5):          # UML dependency: dashed + open head
    arrow(ax, p1, p2, dashed=True, head="->", txt=txt, fs=fs)


def crowfoot(ax, x, y, toward, many=True, optional=False):
    """วาดปลาย relation แบบ crow's foot ที่ (x,y) โดย toward = ทิศเข้าหา entity ('l'/'r')"""
    d = -1 if toward == "r" else 1          # ทิศที่ก้างชี้ออก
    if many:
        for dy in (-1.4, 0, 1.4):
            ax.plot([x, x + d * 3], [y, y + dy], color="#333", lw=1.0)
    else:
        ax.plot([x + d * 2, x + d * 2], [y - 1.4, y + 1.4], color="#333", lw=1.0)
    if optional:
        ax.add_patch(Circle((x + d * 4.5, y), 0.9, fc="white", ec="#333", lw=1.0))


def save(fig, name):
    DIA.mkdir(parents=True, exist_ok=True)
    fig.savefig(DIA / name, dpi=165, bbox_inches="tight")
    plt.close(fig)
    print("  ", (DIA / name).relative_to(BASE))
    return name


# ============ 1. Use Case ============
def d1():
    fig, ax = new(80, 13)
    title(ax, "1. Use Case Diagram", "ระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage)")
    ax.add_patch(Rectangle((21, 4), 62, 66, fc="none", ec="#888", lw=1.2))
    ax.text(52, 67, "ระบบตรวจจับตำหนิพื้นผิวเหล็ก", ha="center", fontsize=8, style="italic")
    op = actor(ax, 8, 36, "ผู้ใช้งาน\n(Operator)")
    eng = actor(ax, 95, 36, "ผู้พัฒนา\n(ML Engineer)")
    L = [oval(ax, 40, y, 24, 7, t) for y, t in [
        (60, "อัปโหลดภาพเหล็ก"), (50, "ปรับ confidence /\nโหมดตรวจละเอียด"),
        (40, "ตรวจภาพ (Stage 1 + Stage 2)"), (30, "ดูผล: กรอบ + ชนิด + ความเสี่ยง"),
        (20, "บันทึก / ส่งออกผล (JSON, ภาพ)")]]
    R = [oval(ax, 66, y, 24, 7, t, CL["data"]) for y, t in [
        (63, "เตรียม dataset"), (53, "เทรนโมเดล Stage 2 (train.py)"), (43, "วัดผล (evaluate.py)"),
        (33, "ปรับ per-class threshold\n(tune_thresholds.py)"), (23, "วัด Stage 1 / วัดกับภาพจริง"),
        (13, "สร้างเอกสาร / ไดอาแกรม")]]
    for u in L:
        arrow(ax, (op["cx"] + 2, u["my"]), (u["l"], u["my"]), head="-")
    for u in R:
        arrow(ax, (eng["cx"] - 2, u["my"]), (u["r"], u["my"]), head="-")
    dep(ax, (L[2]["cx"], L[2]["t"]), (L[0]["cx"], L[0]["b"]), txt="«include»")
    dep(ax, (L[2]["cx"], L[2]["b"]), (L[3]["cx"], L[3]["t"]), txt="«include»")
    return save(fig, "uml_1_usecase.png")


# ============ 2. Context ============
def d2():
    fig, ax = new(84, 12)
    title(ax, "2. Context Diagram (DFD Level 0)")
    c = oval(ax, 50, 42, 26, 15, "ระบบตรวจจับ\nตำหนิพื้นผิวเหล็ก\n(process 0)", CL["sys"], 8.5)
    E = [
        (3, 60, "ผู้ใช้งาน", "ภาพ + พารามิเตอร์  →\n←  ภาพผล + รายการตำหนิ", "r"),
        (3, 14, "ผู้พัฒนา", "dataset + คำสั่ง (CLI)  →\n←  metric, thresholds.json", "r"),
        (73, 62, "Apple DMS46\n(pre-trained)", "DMS46_v1.pt  →", "l"),
        (73, 42, "Ultralytics YOLO", "yolo11s.pt + library  →", "l"),
        (73, 22, "Roboflow\n(ตอนเตรียมข้อมูล)", "NEU / Rust / Crack  →", "l"),
        (37, 4, "Local File System", "↔  weights · thresholds.json · results", "b"),
    ]
    for x, y, nm, lbl, side in E:
        b = box(ax, x, y, 24, 12, nm, CL["ext"], 7.2, round_=False)
        if side == "r":
            arrow(ax, (b["r"], b["my"]), (c["l"] + 2, c["my"]), txt=lbl, fs=5.9)
        elif side == "l":
            arrow(ax, (b["l"], b["my"]), (c["r"] - 2, c["my"]), txt=lbl, fs=5.9)
        else:
            arrow(ax, (b["cx"], b["t"]), (c["cx"], c["b"] + 2), txt=lbl, fs=5.9)
    return save(fig, "uml_2_context.png")


# ============ 3. C4 L1 ============
def d3():
    fig, ax = new(74, 12)
    title(ax, "3. C4 Model — Level 1: System Context")
    op = box(ax, 6, 48, 21, 12, "ผู้ใช้งาน [Person]\nพนักงานคัดกรองเหล็ก", CL["a"], 7.3)
    eng = box(ax, 6, 12, 21, 12, "ผู้พัฒนา [Person]\nดูแลโมเดล/ข้อมูล", CL["a"], 7.3)
    sysb = box(ax, 39, 30, 24, 16,
               "Steel Surface Defect\nDetection System\n[Software System]\n2-stage: DMS46 → YOLO11s", CL["sys"], 7.6)
    dms = box(ax, 76, 50, 20, 11, "Apple DMS46\n[External System]\nmaterial segmentation", CL["ext"], 7.1)
    yolo = box(ax, 76, 31, 20, 11, "Ultralytics YOLO\n[External System]\ndetector runtime", CL["ext"], 7.1)
    fs_ = box(ax, 76, 12, 20, 11, "Local File System\n[External]\nweights / results", CL["ext"], 7.1)
    arrow(ax, (op["r"], op["my"]), (sysb["l"], sysb["my"] + 3), txt="อัปโหลดภาพ / ดูผล\n[HTTP, Gradio]", fs=5.9)
    arrow(ax, (eng["r"], eng["my"]), (sysb["l"], sysb["my"] - 3), txt="เทรน / วัดผล [CLI]", fs=5.9)
    arrow(ax, (sysb["r"], sysb["my"] + 4), (dms["l"], dms["my"]), txt="inference [TorchScript]", fs=5.9)
    arrow(ax, (sysb["r"], sysb["my"]), (yolo["l"], yolo["my"]), txt="inference [Python API]", fs=5.9)
    arrow(ax, (sysb["r"], sysb["my"] - 4), (fs_["l"], fs_["my"]), txt="อ่าน/เขียนไฟล์", fs=5.9)
    return save(fig, "uml_3_c4_l1.png")


# ============ 4. C4 L2 ============
def d4():
    fig, ax = new(92, 12.5)
    title(ax, "4. C4 Model — Level 2: Container")
    ax.add_patch(Rectangle((22, 4), 74, 74, fc="none", ec="#8ca", lw=1.4, ls="--"))
    ax.text(59, 80, "Steel Surface Defect Detection System", ha="center", fontsize=7.6, style="italic")
    op = box(ax, 2, 52, 16, 12, "ผู้ใช้งาน\n[Person]", CL["a"], 7.2)
    eng = box(ax, 2, 16, 16, 12, "ผู้พัฒนา\n[Person]", CL["a"], 7.2)
    ui = box(ax, 26, 56, 24, 12, "Web UI\n[Gradio / Browser]\napp.py — lazy-load models", CL["s2"], 6.9)
    cli = box(ax, 26, 38, 24, 12, "CLI\n[Python]\npipeline.py __main__", CL["s2"], 7.0)
    core = box(ax, 56, 46, 24, 13, "Pipeline Core\n[pipeline.py]\nStage1+2 orchestration,\ncross-region NMS", CL["sys"], 6.8)
    st1 = box(ax, 56, 63, 24, 11, "Stage 1 Runtime\n[PyTorch / TorchScript]\nDMS46 inference", CL["s1"], 6.9)
    st2 = box(ax, 56, 30, 24, 11, "Stage 2 Runtime\n[Ultralytics YOLO]\nYOLO11s inference", CL["s1"], 6.9)
    store = box(ax, 56, 9, 24, 12, "Model & Config Store\n[Local FS]\nDMS46_v1.pt · best.pt ·\nthresholds.json", CL["data"], 6.5)
    scr = box(ax, 26, 9, 24, 12, "Training / Eval Scripts\n[Python]\ntrain · evaluate* ·\ntune_thresholds · fix_labels", CL["data"], 6.5)
    arrow(ax, (op["r"], op["my"]), (ui["l"], ui["my"]), txt="[HTTP :7860]", fs=5.9)
    arrow(ax, (eng["r"], eng["my"]), (cli["l"], cli["my"]), txt="[shell]", fs=5.9)
    dep(ax, (eng["r"], eng["my"] - 3), (scr["l"], scr["my"]))
    dep(ax, (ui["r"], ui["my"]), (core["l"], core["t"] - 2), txt="analyze()")
    dep(ax, (cli["r"], cli["my"]), (core["l"], core["my"]), txt="process_image()")
    arrow(ax, (core["cx"], core["t"]), (st1["cx"], st1["b"]), txt="run_stage1", fs=5.9)
    arrow(ax, (core["cx"], core["b"]), (st2["cx"], st2["t"]), txt="run_stage2", fs=5.9)
    dep(ax, (core["cx"] - 4, core["b"]), (store["cx"] - 4, store["t"]), txt="โหลด weights /\nthresholds.json")
    arrow(ax, (scr["r"], scr["my"]), (store["l"], store["my"]), txt="เขียน best.pt /\nthresholds.json", fs=5.7)
    return save(fig, "uml_4_c4_l2.png")


# ============ 5. ERD ============
def d5():
    fig, ax = new(98, 12.5)
    title(ax, "5. Conceptual Data Model (ERD)",
          "ไม่มีฐานข้อมูล — เอนทิตี persist เป็นไฟล์: *_result.json (RESULT/REGION/DETECTION), "
          "label .txt (LABEL_BOX), real_test/labels.csv (REAL_TEST_LABEL)")

    def ent(x, ytop, name, attrs, fc, wdt=27):
        box(ax, x, ytop - 6, wdt, 6, name, fc, 7.6, round_=False)
        for i, a in enumerate(attrs):
            ax.add_patch(Rectangle((x, ytop - 6 - (i + 1) * 4.4), wdt, 4.4, fc="white", ec="#999", lw=0.6))
            ax.text(x + 1.5, ytop - 6 - i * 4.4 - 2.2, a, fontsize=6.2, va="center")
        bot = ytop - 6 - len(attrs) * 4.4
        return {"cx": x + wdt / 2, "t": ytop, "b": bot, "l": x, "r": x + wdt,
                "hy": ytop - 3, "my": (ytop + bot) / 2}   # hy = กึ่งกลางแถบหัว

    R = ent(3, 90, "RESULT", ["PK  image (path)", "metal_regions : int",
                              "metal_area_ratio : float", "fallback_full_image : bool"], CL["out"])
    RG = ent(3, 56, "REGION   (embedded in RESULT.regions[])",
             ["PK  region_id : int", "box_xywh : int[4]"], CL["s2"], wdt=36)
    D = ent(46, 84, "DETECTION   (REGION.detections[])",
            ["FK  class → DEFECT_CLASS.name", "confidence : float",
             "bbox_xywh : float[4]", "bbox_xyxy_crop : float[4]",
             "bbox_xyxy_global : float[4]", "FK  region_id"], CL["s1"], wdt=34)
    DC = ent(84, 62, "DEFECT_CLASS   (คงที่ในโค้ด)",
             ["PK  id : 0..7", "name : str", "name_th : str", "risk : str"], CL["note"], wdt=15)
    DI = ent(46, 34, "DATASET_IMAGE   (ไฟล์ภาพ)",
             ["PK  filename", "split : {train,valid,test}", "(w, h อ่านตอนโหลด — ไม่เก็บ)"], CL["data"], wdt=34)
    LB = ent(84, 30, "LABEL_BOX  (.txt)",
             ["FK  class_id : 0..7", "xc, yc, w, h : norm"], CL["data"], wdt=15)
    RT = ent(3, 26, "REAL_TEST_LABEL   (labels.csv)",
             ["PK  filename", "classes : set (image-level)"], CL["ext"], wdt=36)

    def rel(pa, pb, la, lb):
        ax.plot([pa[0], pb[0]], [pa[1], pb[1]], color="#444", lw=1.1)
        ax.text(pa[0] + (2.2 if pb[0] > pa[0] else -2.2), pa[1] + 1.6, la, fontsize=7, ha="center", color="#333")
        ax.text(pb[0] + (-2.2 if pb[0] > pa[0] else 2.2), pb[1] + 1.6, lb, fontsize=7, ha="center", color="#333")

    rel((R["cx"], R["b"]), (RG["cx"], RG["t"]), "1", "N")
    rel((RG["r"], RG["hy"]), (D["l"], D["hy"]), "1", "N")
    rel((D["r"], D["hy"]), (DC["l"], DC["hy"]), "N", "1")
    rel((DI["r"], DI["hy"]), (LB["l"], LB["hy"]), "1", "N")
    rel((LB["cx"], LB["t"]), (DC["cx"], DC["b"]), "N", "1")
    ax.text(50, 7, "สัญกรณ์ (Chen): 1 = ความสัมพันธ์ \"หนึ่ง\",  N = \"หลาย\"", ha="center", fontsize=6.6, color="#666")
    return save(fig, "uml_5_erd.png")


# ============ 6. Class ============
def d6():
    fig, ax = new(150, 12)
    ax.text(50, 148, "6. Class Diagram", ha="center", va="top", fontsize=12, fontweight="bold")
    ax.text(50, 143.5, "โค้ด functional — โมดูล = stereotype «utility» (มีแต่ static operation + ค่าคงที่),  "
            "«data» = dict payload,  เส้นประหัวเปิด = dependency", ha="center", va="top",
            fontsize=7.2, color="#666")
    LINE, HEAD = 3.3, 6.0

    def cls(x, ytop, w, stereo, name, ops, attrs=None, fc="#FFFFFF"):
        attrs = attrs or []
        h = HEAD + 2 + LINE * len(attrs) + (2 + LINE * len(ops) if ops else 0)
        ax.add_patch(Rectangle((x, ytop - h), w, h, fc=fc, ec="#555", lw=1.1))
        ax.text(x + w / 2, ytop - 2.2, stereo, ha="center", fontsize=5.8, style="italic")
        ax.text(x + w / 2, ytop - 4.8, name, ha="center", fontsize=7.4, fontweight="bold")
        yy = ytop - HEAD - 1
        ax.plot([x, x + w], [yy, yy], "k-", lw=0.7)
        for a in attrs:
            yy -= LINE
            ax.text(x + 1.4, yy + LINE / 2, a, fontsize=5.9, va="center")
        if ops:
            yy -= 2
            ax.plot([x, x + w], [yy, yy], "k-", lw=0.7)
            for o in ops:
                yy -= LINE
                ax.text(x + 1.4, yy + LINE / 2, o, fontsize=5.9, va="center")
        return {"cx": x + w / 2, "t": ytop, "b": ytop - h, "l": x, "r": x + w, "my": ytop - h / 2}

    P = cls(3, 138, 45, "«utility»", "pipeline", [
        "+ resolve_device(req) : str",
        "+ load_models(dev) : (s1, s2)",
        "+ _preprocess_for_dms(img) : tensor",
        "+ run_stage1(s1, img, dev) : mask",
        "+ mask_to_boxes(mask, ...) : boxes",
        "+ _merge_close_boxes(boxes, gap) : boxes",
        "+ build_regions(mask, shape, min_ratio) : (boxes, meta)",
        "+ run_stage2(s2, crop, conf, dev, augment, class_conf) : det[]",
        "+ _iou_xyxy(a, b) : float",
        "+ cross_region_nms(det[], iou) : det[]",
        "+ load_class_conf(path) : {cls: conf} | None",
        "+ process_image(path, s1, s2, out, conf, dev, ...) : Summary",
        "+ draw_thai_text(img, txt, pos, ...) : img",
        "+ iter_images(folder) / main()"],
        attrs=["DEFECT_CLASSES : str[8]", "DEFECT_INFO : {cls: {name_th, risk}}",
               "METAL_MODEL_INDEX = 22", "STAGE2_MODEL_PATH / THRESHOLDS_PATH : Path"], fc=CL["sys"])
    Ap = cls(3, P["b"] - 6, 45, "«utility»", "app  (Gradio UI)", [
        "+ analyze(img_rgb, conf, detailed) : (img, img, md, rows)",
        "+ build_ui() : gr.Blocks", "- _ensure_models() : (s1, s2, dev)",
        "- _to_bgr(img) : ndarray", "- _stage1_view(img, mask, boxes, meta) : img"],
        attrs=["_STATE : {s1, s2, device, class_conf}", "_RISK_ORDER : {risk: int}"], fc=CL["a"])
    Ev = cls(3, Ap["b"] - 6, 45, "«utility»",
             "evaluate · evaluate_real · evaluate_stage1 · tune_thresholds",
             ["ใช้ pipeline.build_regions / run_stage1 / run_stage2 /", "    cross_region_nms / load_class_conf"], fc=CL["data"])

    De = cls(53, 138, 44, "«data»", "Detection", [], attrs=[
        "class : str", "confidence : float", "region_id : int",
        "bbox_xywh : float[4]", "bbox_xyxy_crop : float[4]", "bbox_xyxy_global : float[4]"], fc=CL["s1"])
    Me = cls(53, De["b"] - 7, 44, "«data»", "RegionMeta", [], attrs=[
        "metal_found : bool", "metal_ratio : float",
        "fallback_full_image : bool", "n_regions : int"], fc=CL["s2"])
    Su = cls(53, Me["b"] - 7, 44, "«data»", "Summary  (*_result.json)", [], attrs=[
        "image : str", "metal_regions : int", "metal_area_ratio : float",
        "fallback_full_image : bool",
        "regions : [{region_id, box_xywh, detections: Detection[]}]"], fc=CL["out"])

    dep(ax, (Ap["l"] + 6, Ap["t"]), (P["l"] + 6, P["b"]), txt="use")     # app  → pipeline
    dep(ax, (Ev["l"] + 16, Ev["t"]), (P["l"] + 16, P["b"]), txt="use")   # eval → pipeline
    dep(ax, (P["r"], De["my"]), (De["l"], De["my"]), txt="«create»")
    dep(ax, (P["r"], Me["my"]), (Me["l"], Me["my"]), txt="«create»")
    dep(ax, (P["r"], Su["my"]), (Su["l"], Su["my"]), txt="«create»")
    return save(fig, "uml_6_class.png")


# ============ 7. Sequence ============
def d7():
    fig, ax = new(104, 13)
    title(ax, "7. Sequence Diagram — analyze() เมื่อผู้ใช้กด \"ตรวจสอบ\" บน UI")
    lanes = [":Operator", ":app (analyze)", ":pipeline", ":DMS46\n(Stage 1)",
             ":YOLO11s\n(Stage 2)", ":File System"]
    xs = [9, 27, 46, 64, 80, 94]
    hdr, bot = 92, 6
    for x, nm in zip(xs, lanes):
        box(ax, x - 7, hdr, 14, 6, nm, CL["a"], 6.6, round_=False)
        ax.plot([x, x], [hdr, bot], color="#999", lw=1, ls=(0, (4, 3)))
    # activation bar ของ pipeline
    ax.add_patch(Rectangle((45, 20), 2, 60, fc="#dfe6ef", ec="#889", lw=0.7))

    def m(i, j, y, t, dashed=False):
        arrow(ax, (xs[i] + (1 if xs[i] < xs[j] else -1), y),
              (xs[j] - (1 if xs[i] < xs[j] else -1), y), dashed=dashed, txt=t, fs=6.0, tx_dy=1.4)

    def note(y, t, x=46, w=52):
        ax.text(x, y, t, fontsize=5.9, ha="center",
                bbox=dict(fc=CL["note"], ec="#999", pad=1.4), zorder=6)

    def frame(y1, y2, x1, x2, lab):
        ax.add_patch(Rectangle((x1, y2), x2 - x1, y1 - y2, fc="none", ec="#c0392b", lw=0.8))
        ax.text(x1 + 4.5, y1 - 1.6, lab, fontsize=5.6, color="#c0392b", ha="center",
                bbox=dict(fc="white", ec="#c0392b", pad=0.7))

    y = 86
    m(0, 1, y, "analyze(image_rgb, conf, detailed)"); y -= 6
    frame(y + 2, y - 12, 30, 99, "opt  [โหลดโมเดลครั้งแรก]")
    m(1, 2, y, "_ensure_models()", dashed=True); y -= 5
    m(2, 5, y, "load DMS46_v1.pt · best.pt · load_class_conf()", dashed=True); y -= 5
    m(5, 2, y, "s1, s2, class_conf", dashed=True); y -= 8
    note(y, "_to_bgr(image_rgb)   (แปลง RGBA/gray → BGR)", x=27, w=30); y -= 6
    m(1, 2, y, "run_stage1(s1, bgr, device)"); y -= 5
    m(2, 3, y, "DMS46 forward", dashed=False); y -= 4
    m(3, 2, y, "metal mask", dashed=True); y -= 6
    note(y, "build_regions(mask, shape)  →  boxes, meta   [metal_ratio < 0.05 → +กรอบทั้งภาพ]\n"
            "_stage1_view(...) → ภาพ Stage 1", x=46, w=60); y -= 9
    frame(y + 3, y - 20, 24, 88, "loop  [แต่ละ region ใน boxes]")
    m(1, 4, y, "run_stage2(s2, crop, conf, device, augment, class_conf)"); y -= 4
    note(y, "ภายใน: [ถ้าโมเดล gray → แปลง crop เป็น gray] ; predict ; กรอง score < class_conf[cls]",
         x=53, w=68); y -= 6
    m(4, 1, y, "detections (bbox → พิกัดภาพเต็ม)", dashed=True); y -= 8
    note(y, "cross_region_nms(flat, iou=0.5)   (ตัดตัวซ้ำข้าม region)", x=27, w=44); y -= 7
    note(y, "วาดกรอบ + ป้ายไทย ; sort ตามความเสี่ยง ; สร้าง verdict", x=27, w=48); y -= 7
    m(1, 0, y, "return (annotated, stage1_img, verdict_md, rows[])", dashed=True)
    return save(fig, "uml_7_sequence.png")


# ============ 8. Activity ============
def d8():
    fig, ax = new(198, 8.5)
    title(ax, "8. Activity Diagram — process_image()  (การประมวลผล 1 ภาพ)")
    ax.add_patch(Circle((50, 190), 2.2, fc="black"))
    y = 188
    prev = (50, 190)

    def step(t, fc, w=60, h=7):
        nonlocal y, prev
        b = box(ax, 50 - w / 2, y - h, w, h, t, CL[fc], 7.0)
        arrow(ax, prev, (b["cx"], b["t"]))
        prev = (b["cx"], b["b"]); y -= h + 5
        return b

    step("image = cv2.imread(path)", "a")
    step("run_stage1(): DMS46 → metal mask", "s1")
    step("build_regions(mask, image.shape, min_metal_ratio)", "s1")
    dg = diamond(ax, 50, y - 8, 46, 16, "metal_ratio < 0.05\nหรือไม่เจอกรอบ?\n(ภายใน build_regions)")
    arrow(ax, prev, (dg["cx"], dg["t"]))
    bf = box(ax, 3, y - 13, 30, 8, "boxes += (0,0,W,H)\nfallback = ทั้งภาพ", CL["note"], 6.7)
    arrow(ax, (dg["l"], dg["my"]), (bf["r"], bf["my"]), txt="ใช่", fs=6.0)
    mrg = y - 20
    arrow(ax, (dg["cx"], dg["b"]), (50, mrg + 1.4), txt="ไม่ใช่", fs=6.0)
    arrow(ax, (bf["cx"], bf["b"]), (50, mrg + 1.4))
    ax.add_patch(Polygon([(50, mrg + 1.6), (52, mrg), (50, mrg - 1.6), (48, mrg)], fc="#333"))
    y = mrg - 6
    prev = (50, mrg - 1.6)
    ax.text(50, y + 2.5, "──  Pass 1: ตรวจทุก region  ──", ha="center", fontsize=6.2, color="#666")
    y -= 3
    step("for region: crop = image[y:y+h, x:x+w]", "s2", w=66)
    step("run_stage2(): [ถ้าโมเดล gray → gray] · predict · กรอง score < class_conf[cls]", "s2", w=78)
    step("แปลง bbox_xyxy_crop → bbox_xyxy_global (บวก offset region)", "s2", w=74)
    step("flat = ทุก detection ;  cross_region_nms(flat, iou=0.5)", "s2", w=66)
    ax.text(50, y + 2.5, "──  Pass 2: วาดผลเฉพาะตัวที่รอด NMS  ──", ha="center", fontsize=6.2, color="#666")
    y -= 3
    dd = diamond(ax, 50, y - 7, 40, 14, "region มี detection\nที่รอด NMS?")
    arrow(ax, prev, (dd["cx"], dd["t"]))
    by = box(ax, 4, y - 21, 42, 8, "วาดกรอบเขียว + กรอบแดง\n+ ป้ายไทย + เก็บลง rows[]", CL["s2"], 6.6)
    bn = box(ax, 54, y - 21, 42, 8, "วาดกรอบเขียว\n+ ป้าย \"เหล็ก #k ปกติ\"", CL["a"], 6.6)
    arrow(ax, (dd["l"], dd["my"]), (by["r"], by["my"]), txt="ใช่", fs=6.0)
    arrow(ax, (dd["r"], dd["my"]), (bn["l"], bn["my"]), txt="ไม่", fs=6.0)
    jn = y - 28
    arrow(ax, (by["cx"], by["b"]), (50, jn + 1.4))
    arrow(ax, (bn["cx"], bn["b"]), (50, jn + 1.4))
    ax.add_patch(Polygon([(50, jn + 1.6), (52, jn), (50, jn - 1.6), (48, jn)], fc="#333"))
    bs = box(ax, 22, jn - 12, 56, 8, "cv2.imwrite(*_result.jpg)\n+ write(*_result.json)", CL["out"], 6.8)
    arrow(ax, (50, jn - 1.6), (bs["cx"], bs["t"]))
    arrow(ax, (bs["cx"], bs["b"]), (50, jn - 18))
    ax.add_patch(Circle((50, jn - 20), 2.4, fc="none", ec="black", lw=1.4))
    ax.add_patch(Circle((50, jn - 20), 1.1, fc="black"))
    return save(fig, "uml_8_activity.png")


# ============ 9. State ============
def d9():
    fig, ax = new(46, 12)
    title(ax, "9. State Machine — วงจรชีวิตของภาพ 1 รูปใน analyze() / process_image()",
          "ระบบไม่มี session ที่คงสถานะข้ามคำขอ — state อยู่ใน local variable เท่านั้น")
    ax.add_patch(Circle((4, 30), 1.6, fc="black"))
    S = []
    xs = [8, 27, 46, 65, 82]
    labs = ["Loaded", "Stage1Done", "RegionsReady", "Stage2Done", "Filtered"]
    subs = ["cv2.imread", "metal mask", "+fallback ถ้า\nratio<0.05", "det[] ต่อ region", "NMS + threshold"]
    ev = ["/ run_stage1", "/ build_regions", "loop / run_stage2", "/ cross_region_nms", None]
    prev = (5.6, 30)
    for i, (x, lab, sub) in enumerate(zip(xs, labs, subs)):
        b = box(ax, x, 24, 15, 11, lab + "\n" + sub, CL["s1"], 6.3)
        arrow(ax, prev, (b["l"], b["my"]), txt=ev[i], fs=5.4, tx_dy=2.3)
        prev = (b["r"], b["my"]); S.append(b)
    fin = box(ax, 38, 6, 26, 8, "Rendered\n(annotated + verdict + rows)", CL["out"], 6.3)
    arrow(ax, (S[-1]["cx"], S[-1]["b"]), (fin["r"], fin["my"]), txt="/ draw + verdict", fs=5.4, tx_dy=2.0)
    arrow(ax, (fin["l"], fin["t"]), (S[2]["cx"], S[2]["b"]), dashed=True,
          txt="[--folder: ภาพถัดไป]", fs=5.4)
    ax.add_patch(Circle((30, 8), 2.0, fc="none", ec="black", lw=1.3))
    ax.add_patch(Circle((30, 8), 0.95, fc="black"))
    arrow(ax, (fin["l"], fin["my"]), (32, 8), txt="[--image: จบ]", fs=5.4)
    return save(fig, "uml_9_state.png")


# ============ 10. Deployment ============
def d10():
    fig, ax = new(88, 12)
    title(ax, "10. Deployment Diagram", "รันบนเครื่องเดียว (ไม่มี server / cloud)")
    ax.add_patch(Rectangle((5, 5), 90, 72, fc="#F7F7F7", ec="#555", lw=1.4))
    ax.text(50, 73, "«device»  Windows 11 Laptop", ha="center", fontsize=8.5, fontweight="bold")
    box(ax, 63, 54, 28, 12, "«device»\nNVIDIA RTX 3050 6GB\nCUDA 12.4", CL["ext"], 7.0)
    ax.add_patch(Rectangle((9, 22), 46, 44, fc="white", ec="#777", lw=1.1))
    ax.text(32, 62, "«execution environment»  Python 3.11 venv", ha="center", fontsize=7.0, fontweight="bold")
    for i, (t, c) in enumerate([
        ("«artifact» pipeline.py · app.py", CL["s2"]),
        ("«artifact» PyTorch 2.6 + Ultralytics 8.4.126", CL["s1"]),
        ("«artifact» DMS46_v1.pt  (Stage 1)", CL["data"]),
        ("«artifact» runs/detect/train-gray-s/weights/best.pt", CL["data"]),
        ("«artifact» thresholds.json", CL["note"])]):
        box(ax, 11, 54 - i * 7, 42, 5.4, t, c, 6.1, round_=False)
    ax.add_patch(Rectangle((60, 26), 30, 20, fc="white", ec="#777", lw=1.1))
    ax.text(75, 42, "«execution environment»  Web Browser", ha="center", fontsize=6.8, fontweight="bold")
    box(ax, 62, 29, 26, 8, "Gradio client\nhttp://127.0.0.1:7860", CL["a"], 6.5, round_=False)
    box(ax, 18, 9, 64, 8, "«file system»  merged_dataset_gray/ · real_test/ · pipeline_results/ · results/",
        CL["ext"], 6.2, round_=False)
    arrow(ax, (74, 29), (33, 52), dashed=True, txt="HTTP :7860", fs=6.0)
    arrow(ax, (32, 22), (42, 17), dashed=True, txt="อ่าน/เขียนไฟล์", fs=6.0)
    arrow(ax, (55, 48), (63, 54), dashed=True, txt="CUDA", fs=6.0)
    return save(fig, "uml_10_deployment.png")


def build_docx(items):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.styles["Normal"].font.name = "Tahoma"
    doc.styles["Normal"].font.size = Pt(11)
    h = doc.add_heading("UML / SA Diagrams — ระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage)", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "เนื้อหาไดอาแกรมอิงโค้ดจริง (pipeline.py / app.py / evaluate*.py / schema ของ *_result.json). "
        "ระบบเป็น ML pipeline + prototype UI (Gradio) — ไม่มีฐานข้อมูล / ระบบล็อกอิน / บัญชีผู้ใช้ จึงใช้ "
        "notation ที่ยังเป็นมาตรฐานแต่ตรงกับของจริง: ERD เป็น Conceptual Data Model (persist เป็นไฟล์), "
        "Class ใช้ stereotype «utility» (คลาสที่มีแต่ static operation) สำหรับโมดูล functional, "
        "State เป็นวงจรชีวิตของภาพ 1 รูปใน local variable. สร้างจาก make_uml_doc.py"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    names = {
        "uml_1_usecase.png": "1. Use Case Diagram",
        "uml_2_context.png": "2. Context Diagram (DFD Level 0)",
        "uml_3_c4_l1.png": "3. C4 Model — Level 1: System Context",
        "uml_4_c4_l2.png": "4. C4 Model — Level 2: Container",
        "uml_5_erd.png": "5. Conceptual Data Model (ERD)",
        "uml_6_class.png": "6. Class Diagram",
        "uml_7_sequence.png": "7. Sequence Diagram — analyze()",
        "uml_8_activity.png": "8. Activity Diagram — process_image()",
        "uml_9_state.png": "9. State Machine — วงจรชีวิตของภาพ",
        "uml_10_deployment.png": "10. Deployment Diagram",
    }
    desc = {
        "uml_1_usecase.png": "Actor 2 ราย. ผู้ใช้งาน: อัปโหลดภาพ, ปรับ confidence/โหมดตรวจละเอียด, ตรวจภาพ, ดูผล, "
        "บันทึก/ส่งออก. ผู้พัฒนา: เตรียม dataset, เทรน (train.py), วัดผล (evaluate.py), ปรับ threshold "
        "(tune_thresholds.py), วัด Stage 1 / ภาพจริง, สร้างเอกสาร. \"ตรวจภาพ\" «include» การอัปโหลดและการดูผล.",
        "uml_2_context.png": "process 0 = ทั้งระบบ. รับภาพ+พารามิเตอร์จากผู้ใช้ผ่าน UI คืนภาพผล+รายการตำหนิ; รับ "
        "dataset/คำสั่ง CLI จากผู้พัฒนา คืน metric/thresholds.json. External: ไฟล์โมเดล DMS46, library YOLO, "
        "ระบบไฟล์ท้องถิ่น. Roboflow เกี่ยวเฉพาะตอนเตรียมข้อมูล (offline, ครั้งเดียว).",
        "uml_3_c4_l1.png": "มุมมอง C4 สูงสุด: ผู้ใช้ (HTTP/Gradio) และผู้พัฒนา (CLI) ↔ ระบบ ↔ ระบบภายนอก 3 ตัว "
        "(DMS46 TorchScript, Ultralytics YOLO runtime, file system). ไม่มีการเรียก API ผ่านเครือข่าย.",
        "uml_4_c4_l2.png": "Container ภายในระบบ: Web UI (Gradio, lazy-load model), CLI, Pipeline Core (pipeline.py — "
        "orchestrate + cross-region NMS), Stage 1/2 Runtime, Model & Config Store (ไฟล์บนดิสก์), Training/Eval Scripts. "
        "เส้นทึบ = การเรียก inference, เส้นประ = dependency/อ่าน-เขียนไฟล์. ไม่มี container ฐานข้อมูล.",
        "uml_5_erd.png": "Conceptual Data Model — ไม่มี DB. RESULT (1) ── (N) REGION (embedded ใน regions[]) "
        "(1) ── (N) DETECTION (N) ── (1) DEFECT_CLASS (คงที่ในโค้ด: DEFECT_CLASSES + DEFECT_INFO). ฝั่งข้อมูลเทรน "
        "DATASET_IMAGE (1) ── (N) LABEL_BOX (.txt YOLO) (N) ── (1) DEFECT_CLASS. REAL_TEST_LABEL (labels.csv) เป็น "
        "label ระดับภาพ (multi-label) สำหรับ evaluate_real.py. สัญกรณ์ตีนกา (crow's foot).",
        "uml_6_class.png": "โค้ด functional — โมดูล pipeline / app / evaluate* เป็น «utility» (มีแต่ static operation "
        "+ ค่าคงที่ระดับโมดูล). «data» Detection / RegionMeta / Summary คือ dict ที่ pipeline สร้างและส่งต่อ "
        "(Summary = โครงของ *_result.json). app และ evaluate* «use» pipeline; pipeline «create» dict payload ทั้งสาม.",
        "uml_7_sequence.png": "ลำดับใน app.analyze(): opt โหลดโมเดลครั้งแรก (_ensure_models → load_models + "
        "load_class_conf) → _to_bgr → run_stage1 → build_regions (+fallback) + _stage1_view → loop run_stage2 ต่อ region "
        "(แปลง gray + กรอง per-class threshold ภายในฟังก์ชัน) → cross_region_nms → วาดผล + verdict → return 4 ค่า "
        "(annotated, stage1_img, verdict_md, rows[]).",
        "uml_8_activity.png": "Flow ของ process_image(): อ่านภาพ → run_stage1 → build_regions (จุดตัดสินใจ "
        "metal_ratio < 0.05 → เพิ่มกรอบทั้งภาพ อยู่ภายในฟังก์ชันนี้) → Pass 1: loop crop + run_stage2 "
        "(gray + threshold ภายใน) + แปลง bbox เป็นพิกัดภาพเต็ม → cross_region_nms → Pass 2: จุดตัดสินใจต่อ region "
        "(มี detection ที่รอด NMS?) วาดผล/ทำเครื่องหมายปกติ → เขียน .jpg + .json.",
        "uml_9_state.png": "state ของ \"ภาพ 1 รูป\" ระหว่างประมวลผล (อยู่ใน local variable ไม่ persist): Loaded → "
        "Stage1Done → RegionsReady (+fallback ถ้า ratio<0.05) → Stage2Done → Filtered (หลัง NMS+threshold) → Rendered. "
        "transition กำกับด้วยการเรียกฟังก์ชัน. โหมด --folder วนกลับไป RegionsReady สำหรับภาพถัดไป.",
        "uml_10_deployment.png": "โน้ตบุ๊ก Windows 11 เครื่องเดียว: Python 3.11 venv (โค้ด + PyTorch 2.6/Ultralytics "
        "8.4.126 + DMS46_v1.pt + train-gray-s/best.pt + thresholds.json), เบราว์เซอร์ต่อ Gradio ที่ 127.0.0.1:7860, "
        "GPU RTX 3050 ผ่าน CUDA 12.4, ข้อมูล/ผลลัพธ์บนระบบไฟล์ท้องถิ่น. ไม่มี server / cloud.",
    }
    for img in items:
        doc.add_page_break()
        doc.add_heading(names[img], level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(DIA / img), width=Inches(6.6))
        doc.add_paragraph(desc[img])
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX))
    print("\nเขียน:", DOCX.relative_to(BASE))


if __name__ == "__main__":
    print("สร้าง UML/SA diagrams -> figures/diagrams/")
    items = [d1(), d2(), d3(), d4(), d5(), d6(), d7(), d8(), d9(), d10()]
    build_docx(items)
