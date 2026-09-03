"""
สร้างเอกสาร Word รวมไดอาแกรมทั้งหมดของระบบ -> docs/system_diagrams.docx
+ รูป PNG แยกไฟล์ใน figures/diagrams/

    python make_diagrams_doc.py

ต้องมี: python-docx, matplotlib  (pip install python-docx)
ไดอาแกรม:
  1. สถาปัตยกรรม 2-Stage (ภาพรวม)
  2. Stage 1 — DMS46 Metal Localization (ภายใน)
  3. Stage 2 — YOLO11 Defect Detection (ภายใน)
  4. การเตรียมข้อมูล (dataset pipeline)
  5. โครงการวัดผล (evaluation)
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib import font_manager

BASE = Path(__file__).resolve().parent
DIA = BASE / "figures" / "diagrams"
DOCX = BASE / "docs" / "system_diagrams.docx"

# ---- ฟอนต์ไทย (Tahoma มีทั้งไทย + ลูกศร) ----
for _n in ("Tahoma", "Leelawadee UI", "Angsana New"):
    try:
        font_manager.findfont(_n, fallback_to_default=False)
        plt.rcParams["font.family"] = _n
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False

C = {  # สีตามหน้าที่
    "in": "#E8EAF0", "s1": "#CFE2F3", "s2": "#D9EAD3",
    "data": "#FCE5CD", "out": "#EAD1DC", "note": "#FFF2CC",
}
BOX_W, BOX_H, GAP = 7.6, 1.35, 0.75
X = 1.2


def _box(ax, cx, y_top, text, fc, w=BOX_W, h=BOX_H, fs=9):
    """วาดกล่องที่ (จุดกึ่งกลาง x = cx, ขอบบน y = y_top). คืน (cx, y_bottom)"""
    ax.add_patch(FancyBboxPatch((cx - w / 2, y_top - h), w, h,
                                boxstyle="round,pad=0.02,rounding_size=0.05",
                                fc=fc, ec="#555555", lw=1.1))
    ax.text(cx, y_top - h / 2, text, ha="center", va="center", fontsize=fs)
    return cx, y_top - h


def _arrow(ax, p_from, p_to):
    ax.add_patch(FancyArrowPatch(p_from, p_to, arrowstyle="-|>", mutation_scale=15,
                                 lw=1.3, color="#3b3b3b"))


def _vstack(ax, steps, x=X, w=BOX_W, h=BOX_H, gap=GAP, y_start=None):
    """steps = [(fc, text, fontsize?), ...] เรียงจากบนลงล่าง"""
    n = len(steps)
    total = n * h + (n - 1) * gap
    y = total if y_start is None else y_start
    prev = None
    for st in steps:
        fc, txt = st[0], st[1]
        fs = st[2] if len(st) > 2 else 9
        cx = x + w / 2
        if prev:
            _arrow(ax, prev, (cx, y))
        prev = _box(ax, cx, y, txt, C[fc], w, h, fs)
        y -= h + gap
    return total


def _fig(content_h, width_in=8.5):
    h_in = max(4.0, content_h * 0.62)
    fig, ax = plt.subplots(figsize=(width_in, h_in))
    ax.set_xlim(0, 10)
    ax.set_ylim(-0.3, content_h + 1.4)
    ax.axis("off")
    return fig, ax


def _title(ax, content_h, text, fs=12):
    ax.text(5, content_h + 0.7, text, ha="center", va="center",
            fontsize=fs, fontweight="bold")


def _save(fig, name):
    DIA.mkdir(parents=True, exist_ok=True)
    p = DIA / name
    fig.savefig(p, dpi=170, bbox_inches="tight")
    plt.close(fig)
    print("  ", p.relative_to(BASE))
    return p


# ---------------- Diagram 1 : ภาพรวม ----------------
def d1():
    steps = [
        ("in", "ภาพถ่ายเหล็ก (BGR)  —  pipeline.py / app.py"),
        ("s1", "STAGE 1 — DMS46 (TorchScript)\nหา region ที่เป็นวัสดุ \"Metal\""),
        ("s1", "mask พื้นที่เหล็ก (0/255)\nmask_to_boxes(): contour + รวมกรอบ + กรอง"),
        ("note", "fallback: metal_ratio < 0.05  ->  เพิ่มกรอบ = ทั้งภาพ\n(build_regions)"),
        ("s2", "STAGE 2 — YOLO11s (เทรนเอง, grayscale)\nรันตรวจตำหนิบนแต่ละ region"),
        ("s2", "map bbox -> พิกัดภาพเต็ม  ->  cross-region NMS\n(class-aware, IoU 0.5)"),
        ("s2", "per-class confidence threshold\n(thresholds.json / tune_thresholds.py)"),
        ("out", "ผลลัพธ์: ภาพ + กรอบ + ป้ายไทย + ความเสี่ยง\n*_result.jpg  /  *_result.json"),
    ]
    total = 8 * BOX_H + 7 * GAP
    fig, ax = _fig(total)
    _vstack(ax, steps)
    _title(ax, total, "ไดอาแกรม 1 — สถาปัตยกรรม 2-Stage (ภาพรวม)")
    return _save(fig, "d1_architecture.png")


# ---------------- Diagram 2 : Stage 1 ----------------
def d2():
    steps = [
        ("in", "crop / ภาพเต็ม (BGR)", 9),
        ("s1", "_preprocess_for_dms(): resize ด้านยาว = 512 (รักษาสัดส่วน)\n+ ImageNet normalize (สเกล 0–255)", 8.5),
        ("s1", "DMS46 forward  ->  label map [1,1,H,W]  (46 วัสดุ, argmax ในตัวโมเดล)", 8.5),
        ("s1", "เลือก output index == 22 (\"Metal\")  ->  binary mask\nresize กลับขนาดเดิม (NEAREST)", 8.5),
        ("s1", "morphology OPEN + CLOSE (5x5)  ->  cv2.findContours (RETR_EXTERNAL)", 8.5),
        ("s1", "กรองกรอบ: area >= 500, fill_ratio >= 0.10, ไม่ใหญ่เกิน 98%\n_merge_close_boxes(gap ~ 3% ของด้านสั้น)", 8.5),
        ("s1", "ทิ้งกรอบจิ๋ว (<1%) ถ้ามีกรอบใหญ่ (>5%) อยู่แล้ว", 8.5),
        ("note", "build_regions(): ถ้า metal_ratio < 0.05 หรือไม่เจอกรอบ\n->  ต่อกรอบ (0, 0, W, H) เป็น fallback", 8.5),
        ("out", "list ของ (x, y, w, h) ในระบบพิกัดภาพเต็ม  +  meta", 9),
    ]
    total = 9 * BOX_H + 8 * GAP
    fig, ax = _fig(total)
    _vstack(ax, steps)
    _title(ax, total, "ไดอาแกรม 2 — Stage 1: DMS46 Metal Localization (ภายใน)", 11)
    ax.text(5, total + 0.15, "pipeline.py : run_stage1 / mask_to_boxes / _merge_close_boxes / build_regions",
            ha="center", fontsize=8.5, color="#555")
    return _save(fig, "d2_stage1.png")


# ---------------- Diagram 3 : Stage 2 ----------------
def d3():
    steps = [
        ("in", "regions จาก Stage 1  (รวมกรอบ fallback ถ้ามี)", 9),
        ("s2", "สำหรับแต่ละ region: crop ภาพ", 9),
        ("s2", "ถ้าโมเดลเทรนบน grayscale  ->  แปลง crop เป็น gray-3ch\n(ตรวจอัตโนมัติจาก ckpt: train_args.data มีคำว่า 'gray')", 8.5),
        ("s2", "YOLO11s.predict(conf = min(--conf, per-class))\naugment=True ได้ (TTA — โหมด \"ตรวจละเอียด\" ใน app)", 8.5),
        ("s2", "detection  ->  บวก offset ของ region  ->  bbox_xyxy_global", 8.5),
        ("s2", "รวม detection ทุก region  ->  cross_region_nms()\nตัดตัวซ้ำ (คลาสเดียวกัน, IoU_global > 0.5)", 8.5),
        ("s2", "กรองด้วย per-class threshold (thresholds.json)\nเช่น rust >= 0.90, crazing >= 0.59, crack >= 0.36", 8.5),
        ("out", "detections สุดท้าย  ->  วาดกรอบ + ป้ายไทย + ตารางความเสี่ยง", 9),
    ]
    total = 8 * BOX_H + 7 * GAP
    fig, ax = _fig(total)
    _vstack(ax, steps)
    _title(ax, total, "ไดอาแกรม 3 — Stage 2: YOLO11 Defect Detection (ภายใน)", 11)
    ax.text(5, total + 0.15, "pipeline.py : run_stage2 / cross_region_nms / load_class_conf",
            ha="center", fontsize=8.5, color="#555")
    return _save(fig, "d3_stage2.png")


# ---------------- Diagram 4 : การเตรียมข้อมูล ----------------
def d4():
    steps = [
        ("data", "merge_datasets.py — รวม 3 dataset เป็น 8 คลาส\nกรองไฟล์ปนเปื้อนด้วย prefix ชื่อไฟล์ + remap class id + polygon->bbox", 8.3),
        ("data", "resplit_dataset.py — แบ่ง train/valid/test ใหม่ (stratified)\nทุก split ครบ 8 คลาส   (3353 / 416 / 416)", 8.3),
        ("note", "fix_labels.py — รวมกล่อง crazing(0) / rolled-in_scale(4) เป็น union 1 กล่อง/ภาพ\n+ ตัดกล่อง degenerate   (สำรองไว้ที่ labels_raw/)", 8.3),
        ("note", "make_grayscale_dataset.py  ->  merged_dataset_gray/\ngrayscale 3 ช่อง (ตัด shortcut ที่โมเดลใช้ \"สี\" แยกโดเมน)", 8.3),
        ("data", "make_oversampled_list.py --dataset merged_dataset_gray\ncrazing x3 ; pitted / rolled-in / scratches x2  ->  train_oversampled.txt", 8.3),
        ("s2", "train.py --recipe texture --model yolo11s.pt\n--data merged_dataset_gray/data_oversampled.yaml   ->   run: train-gray-s", 8.3),
    ]
    total = 6 * BOX_H + 5 * GAP + 2.0            # เผื่อแถวแหล่งข้อมูลด้านบน
    fig, ax = _fig(total, width_in=9.5)
    src_y = total + 0.0
    src_w = 2.9
    srcs = [
        (1.9, "NEU-DET (Roboflow)\n6 คลาส texture\n(ปนภาพ Rust ~932)"),
        (5.0, "Rust dataset\nRUST + DANGER-RUST\n(ตัด NO-RUST)"),
        (8.1, "Crack dataset\n(polygon -> bbox)\ncap 800 / 150 / 150"),
    ]
    first_cx = X + BOX_W / 2
    for cx, txt in srcs:
        _box(ax, cx, src_y, txt, C["data"], w=src_w, h=1.5, fs=7.8)
        _arrow(ax, (cx, src_y - 1.5), (first_cx, total - 2.0 + 0.0 + BOX_H))
    _vstack(ax, steps, y_start=total - 2.0)
    _title(ax, total, "ไดอาแกรม 4 — การเตรียมข้อมูล (raw  ->  merge  ->  resplit  ->  fix  ->  gray  ->  oversample  ->  train)", 10.5)
    return _save(fig, "d4_data_prep.png")


# ---------------- Diagram 5 : โครงการวัดผล ----------------
def d5():
    total = 8.0
    fig, ax = _fig(total, width_in=10)
    cx_c, cy_c = 5.0, total
    _box(ax, cx_c, cy_c, "โมเดล / ระบบ ที่จะวัดผล", C["out"], w=3.6, h=1.2, fs=9)
    leaves = [
        (1.7, 5.6, "s2", "evaluate.py --mode stage2\nmAP50 / mAP50-95 ต่อคลาส\n(test split, ผ่าน ultralytics)"),
        (5.0, 5.6, "s2", "tune_thresholds.py\nF1-vs-conf ต่อคลาส (val)\n->  thresholds.json"),
        (8.3, 5.6, "s1", "evaluate_stage1.py\nfallback rate / coverage /\ngt_area_kept / เวลา"),
        (3.3, 2.7, "out", "evaluate.py --mode pipeline\nend-to-end image-level P/R/F1\n(crop NEU — Stage 1 แทบไม่มีผล)"),
        (7.0, 2.7, "note", "evaluate_real.py   [ต้องทำ]\nreal_test/ ภาพถ่ายจริง\npipeline  vs  baseline (ไม่มี Stage 1)"),
    ]
    for cx, cy, fc, txt in leaves:
        _box(ax, cx, cy, txt, C[fc], w=3.0, h=1.7, fs=7.6)
        _arrow(ax, (cx_c, cy_c - 1.2), (cx, cy))
    ax.text(5, 0.7, "[ต้องทำ] evaluate_real.py — เป็นการทดลองเดียวที่พิสูจน์คุณค่าของ Stage 1 (ยังไม่ได้เก็บภาพ)",
            ha="center", fontsize=8.4, color="#a00000")
    _title(ax, total, "ไดอาแกรม 5 — โครงการวัดผล (แต่ละสคริปต์วัดอะไร)", 11)
    return _save(fig, "d5_eval.png")


# ---------------- ประกอบเป็น .docx ----------------
def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Tahoma"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("ไดอาแกรมระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "เอกสารรวมไดอาแกรมสถาปัตยกรรมและกระบวนการทั้งหมดของระบบ",
        "สร้างจาก make_diagrams_doc.py — แก้เนื้อหาที่สคริปต์แล้วรันใหม่เพื่อ regenerate",
        "Stage 1: DMS46 (Apple Dense Material Segmentation, pre-trained TorchScript)   |   "
        "Stage 2: YOLO11s เทรนเองบน dataset รวม 8 คลาส (grayscale, run: train-gray-s)",
    ):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("1. สถาปัตยกรรม 2-Stage (ภาพรวม)", "d1_architecture.png",
         "ภาพถ่ายเข้าระบบ -> Stage 1 (DMS46) หา region ที่เป็นโลหะแล้วสร้าง mask -> แปลง mask เป็นกรอบ; "
         "ถ้าเจอโลหะน้อยกว่า 5% ของภาพจะเพิ่มกรอบ \"ทั้งภาพ\" เป็น fallback -> Stage 2 (YOLO11s) "
         "ตรวจชนิดตำหนิบนแต่ละกรอบ -> รวมผลข้ามกรอบด้วย cross-region NMS -> กรองด้วย confidence "
         "threshold รายคลาส -> วาดผลและบันทึก JSON. ไฟล์หลัก: pipeline.py (CLI), app.py (UI Gradio)."),
        ("2. Stage 1 — DMS46 Metal Localization (ภายใน)", "d2_stage1.png",
         "preprocess (resize ด้านยาว 512 + ImageNet normalize สเกล 0–255 ให้ตรงกับ ml-dms-dataset ต้นฉบับ) "
         "-> DMS46 คืน label map 46 วัสดุ -> เลือก output index 22 (\"Metal\") -> binary mask -> "
         "morphology + findContours -> กรองกรอบด้วยพื้นที่/สัดส่วนการเติม -> รวมกรอบที่อยู่ติดกัน -> "
         "fallback ทั้งภาพเมื่อ metal_ratio < 0.05. ฟังก์ชัน: run_stage1(), mask_to_boxes(), "
         "_merge_close_boxes(), build_regions()."),
        ("3. Stage 2 — YOLO11 Defect Detection (ภายใน)", "d3_stage2.png",
         "แต่ละ region ถูก crop แล้วส่งเข้า YOLO11s. ถ้าโมเดลเทรนบน grayscale (ตรวจอัตโนมัติจาก checkpoint) "
         "จะแปลง crop เป็น grayscale ก่อนเพื่อไม่ให้เกิด train/serve skew. จากนั้นแปลง bounding box กลับเป็น "
         "พิกัดภาพเต็ม รวมทุก region แล้วทำ class-aware NMS ข้าม region (ตัด detection ซ้ำจากกรอบที่ทับกัน "
         "หรือจาก fallback) และกรองด้วย threshold รายคลาสจาก thresholds.json (สร้างโดย tune_thresholds.py). "
         "ฟังก์ชัน: run_stage2(), cross_region_nms(), load_class_conf()."),
        ("4. การเตรียมข้อมูล (Dataset Pipeline)", "d4_data_prep.png",
         "รวม 3 dataset (NEU-DET, Rust, Crack) เป็น 8 คลาส โดย merge_datasets.py กรองไฟล์ที่ปนเปื้อน "
         "และแปลง polygon -> bbox -> resplit_dataset.py แบ่ง split ใหม่แบบ stratified ให้ทุก split ครบ 8 คลาส "
         "-> fix_labels.py รวมกล่อง crazing / rolled-in_scale เป็นกล่องเดียวต่อภาพ (ทั้งสองเป็น texture ทั้ง patch) "
         "และตัดกล่องเสีย -> make_grayscale_dataset.py สร้างเวอร์ชัน grayscale เพื่อตัด shortcut เรื่องสี "
         "-> make_oversampled_list.py ทำ class-balanced oversampling -> train.py (yolo11s, recipe texture). "
         "ขั้นตอนละเอียดใน prepare_data.md."),
        ("5. โครงการวัดผล (Evaluation)", "d5_eval.png",
         "evaluate.py --mode stage2 วัด mAP ของ Stage 2 บน test split; tune_thresholds.py หา conf ที่ดีที่สุด "
         "รายคลาสจาก val; evaluate_stage1.py วัดว่า DMS46 ทำงานได้แค่ไหนเชิงตัวเลข; evaluate.py --mode pipeline "
         "วัด end-to-end แบบ image-level (แต่บน crop NEU ซึ่ง Stage 1 แทบไม่มีผล); evaluate_real.py วัดบนภาพถ่ายจริง "
         "(real_test/) เทียบ pipeline ที่มี Stage 1 กับ baseline ที่เป็น YOLO บนภาพเต็ม — เป็นการทดลองเดียวที่พิสูจน์ "
         "คุณค่าของ Stage 1 ได้."),
    ]
    for title, img, desc in sections:
        doc.add_page_break()
        doc.add_heading(title, level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(DIA / img), width=Inches(6.4))
        doc.add_paragraph(desc)

    doc.add_page_break()
    doc.add_heading("ภาคผนวก — คลาสตำหนิ 8 ประเภท", level=1)
    rows = [
        ("id", "class", "ชื่อไทย", "ที่มา", "ความเสี่ยง"),
        ("0", "crazing", "รอยแตกลายงา", "NEU", "ปานกลาง–สูง"),
        ("1", "inclusion", "สิ่งแปลกปลอมฝังใน", "NEU", "ปานกลาง"),
        ("2", "patches", "รอยแผ่น/ผิวลอก", "NEU", "ต่ำ–ปานกลาง"),
        ("3", "pitted_surface", "ผิวขรุขระเป็นหลุม", "NEU", "ปานกลาง"),
        ("4", "rolled-in_scale", "สะเก็ดฝังจากการรีด", "NEU", "ปานกลาง"),
        ("5", "scratches", "รอยขีดข่วน", "NEU", "ต่ำ"),
        ("6", "rust", "สนิม", "Rust dataset (RUST + DANGER-RUST)", "สูง"),
        ("7", "crack", "รอยแตกร้าว", "Crack dataset", "สูง"),
    ]
    tb = doc.add_table(rows=len(rows), cols=5)
    try:
        tb.style = "Light Grid Accent 1"
    except KeyError:
        tb.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            tb.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph("รูป PNG แต่ละไดอาแกรมอยู่ใน figures/diagrams/ — regenerate ทั้งหมดด้วย  "
                      "python make_diagrams_doc.py")

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX))
    print("\nเขียน:", DOCX.relative_to(BASE))


if __name__ == "__main__":
    print("สร้างไดอาแกรม -> figures/diagrams/")
    d1(); d2(); d3(); d4(); d5()
    build_docx()
