"""
สร้างเอกสารออกแบบระบบ -> docs/design_document.docx
+ รูป wireframe -> figures/diagrams/ui_wireframe.png

    python make_design_doc.py

หัวข้อ: Functional / Non-functional Requirements, System Architecture,
        Database (Data) Design, UI/UX Design, Flowchart / Use Case
รูปสถาปัตยกรรม/UML ดึงจากที่ make_diagrams_doc.py และ make_uml_doc.py สร้างไว้
(รันสองไฟล์นั้นก่อนถ้ายังไม่มีรูป)

ต้องมี: python-docx, matplotlib
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle, Circle
from matplotlib import font_manager

BASE = Path(__file__).resolve().parent
DIA = BASE / "figures" / "diagrams"
DOCX = BASE / "docs" / "design_document.docx"

for _n in ("Tahoma", "Leelawadee UI", "Angsana New"):
    try:
        font_manager.findfont(_n, fallback_to_default=False)
        plt.rcParams["font.family"] = _n
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False


# ---------- UI wireframe ----------
def make_wireframe():
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 74)
    ax.axis("off")

    def panel(x, y, w, h, label, fc="#F4F6FA", sub=None):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.4,rounding_size=1",
                                    fc=fc, ec="#8892a6", lw=1.1))
        ax.text(x + 1.5, y + h - 2.2, label, fontsize=7.4, fontweight="bold", va="top")
        if sub:
            ax.text(x + w / 2, y + h / 2 - 1, sub, fontsize=6.6, ha="center", va="center", color="#666")

    ax.text(50, 72, "UI Wireframe — app.py  (Gradio Blocks, http://127.0.0.1:7860)",
            ha="center", fontsize=11, fontweight="bold")

    # header
    panel(3, 62, 94, 7, "gr.Markdown  (หัวเรื่อง)",
          sub="ตรวจจับตำหนิพื้นผิวเหล็ก — Prototype   |   Stage 1 DMS46 → Stage 2 YOLO11s (8 คลาส)   |   \"ผู้ช่วยคัดกรอง ไม่ใช่ระบบใช้งานจริง\"")

    # row 1 : input column + output column
    panel(3, 30, 46, 30, "gr.Row  →  gr.Column (ซ้าย: input)", fc="#EEF3FF")
    panel(6, 47, 40, 10.5, "gr.Image  \"ภาพเหล็กที่จะตรวจ\"  (อัปโหลด/วาง)", fc="#FFFFFF")
    panel(6, 42, 40, 3.6, "gr.Slider 0.1–0.9  \"Confidence threshold (Stage 2)\"  (ค่าเริ่ม 0.4)", fc="#FFFFFF")
    panel(6, 38, 40, 3.4, "gr.Checkbox  \"ตรวจละเอียด (TTA — ช้าลง ~2–3x)\"", fc="#FFFFFF")
    ax.add_patch(FancyBboxPatch((6, 34.4), 15, 3.0, boxstyle="round,pad=0.3,rounding_size=1",
                                fc="#c0392b", ec="none"))
    ax.text(13.5, 35.9, "ตรวจสอบ", ha="center", va="center", fontsize=7, color="white", fontweight="bold")
    panel(23, 34.2, 23, 3.4, "gr.Examples  (test_images/*)", fc="#FFFFFF")

    panel(51, 30, 46, 30, "gr.Column (ขวา: output)", fc="#EEFBF0")
    panel(54, 47, 40, 10.5, "gr.Markdown  verdict",
          sub="\"พบตำหนิ N ชนิด\" / \"ความเสี่ยงสูง: ...\" / \"ไม่พบตำหนิ\"  +  Stage1 %, Stage2 จุด, device", fc="#FFFFFF")
    panel(54, 32, 40, 13, "gr.Dataframe  \"รายการตำหนิ (เรียงตามความเสี่ยง)\"", fc="#FFFFFF",
          sub="บริเวณ | class | ชนิด(ไทย) | ความมั่นใจ | ความเสี่ยง")

    # row 2 : two result images
    panel(3, 4, 46, 24, "gr.Row → gr.Image  \"ผลลัพธ์\"", fc="#FFF9F0",
          sub="ภาพต้นฉบับ + กรอบเขียว (region เหล็ก) + กรอบแดง (ตำหนิ) + ป้ายไทย")
    panel(51, 4, 46, 24, "gr.Image  \"Stage 1 — พื้นที่ที่เป็นเหล็ก\"", fc="#FFF9F0",
          sub="overlay เขียว = เหล็ก, ส้ม = กรอบ fallback ทั้งภาพ")

    ax.text(50, 1.4, "btn.click(analyze, inputs=[image, conf, detailed], "
            "outputs=[result_img, stage1_img, verdict, table])",
            ha="center", fontsize=6.6, color="#444", style="italic")

    DIA.mkdir(parents=True, exist_ok=True)
    fig.savefig(DIA / "ui_wireframe.png", dpi=165, bbox_inches="tight")
    plt.close(fig)
    print("  ", (DIA / "ui_wireframe.png").relative_to(BASE))


# ---------- docx ----------
FR = [
    ("FR-01", "รับภาพนำเข้า", "ผู้ใช้อัปโหลดภาพเหล็ก (jpg/jpeg/png/bmp/webp/tif) ผ่าน UI หรือระบุ --image / --folder ผ่าน CLI", "สูง"),
    ("FR-02", "Stage 1 — Metal Localization", "ตรวจหาพื้นที่ที่เป็นวัสดุโลหะด้วย DMS46 แล้วแปลงเป็น bounding regions (mask → contour → merge)", "สูง"),
    ("FR-03", "Fallback ตรวจทั้งภาพ", "ถ้าพื้นที่โลหะ < 5% ของภาพ หรือไม่พบ region ให้เพิ่มกรอบ = ทั้งภาพ เพื่อไม่พลาดตำหนิ", "สูง"),
    ("FR-04", "Stage 2 — Defect Detection", "ตรวจจับและจำแนกตำหนิ 8 ประเภท (crazing, inclusion, patches, pitted_surface, rolled-in_scale, scratches, rust, crack) บนแต่ละ region ด้วย YOLO11s", "สูง"),
    ("FR-05", "ปรับ Confidence threshold", "ผู้ใช้ปรับ threshold รวมได้ (slider 0.1–0.9); ระบบใช้ per-class threshold จาก thresholds.json อัตโนมัติถ้ามีไฟล์", "กลาง"),
    ("FR-06", "โหมดตรวจละเอียด (TTA)", "เปิด test-time augmentation ของ Ultralytics เพื่อเพิ่ม recall (แลกกับความเร็ว ~2–3x)", "ต่ำ"),
    ("FR-07", "Cross-region NMS", "ตัด detection ที่ซ้ำกันข้าม region (คลาสเดียวกัน, IoU ในระบบพิกัดภาพเต็ม > 0.5)", "กลาง"),
    ("FR-08", "แสดงผลบน UI", "วาดกรอบ region (เขียว) + กรอบตำหนิ (แดง) + ป้ายภาษาไทย + ระดับความเสี่ยง; ตารางรายการตำหนิเรียงตามความเสี่ยง; ภาพผล Stage 1; สรุปผล (verdict)", "สูง"),
    ("FR-09", "บันทึกผลลัพธ์ (CLI)", "เขียน <name>_result.jpg และ <name>_result.json ต่อภาพ; โหมด --folder สร้าง _index.json สรุปทั้งชุด", "สูง"),
    ("FR-10", "รองรับภาพหลายรูปแบบ", "แปลง grayscale / RGBA เป็น BGR; ถ้าโมเดล Stage 2 เทรนบน grayscale ระบบแปลง crop เป็น grayscale ก่อนตรวจอัตโนมัติ", "กลาง"),
    ("FR-11", "เตรียมข้อมูล (ผู้พัฒนา)", "merge 3 dataset → resplit stratified → fix_labels → grayscale → oversample (merge_datasets/resplit_dataset/fix_labels/make_grayscale_dataset/make_oversampled_list)", "สูง"),
    ("FR-12", "เทรนโมเดล Stage 2 (ผู้พัฒนา)", "train.py — YOLO11n/s, recipe {default, texture}, resume ได้, validation บน test split อัตโนมัติ", "สูง"),
    ("FR-13", "วัดผล (ผู้พัฒนา)", "evaluate.py (mAP รายคลาส / pipeline image-level), evaluate_stage1.py (Stage 1 เชิงตัวเลข), evaluate_real.py (ภาพจริง: pipeline vs baseline), tune_thresholds.py (per-class conf)", "สูง"),
    ("FR-14", "สร้างรูป/เอกสารประกอบ (ผู้พัฒนา)", "make_figures.py, make_diagrams_doc.py, make_uml_doc.py, make_design_doc.py", "ต่ำ"),
]

NFR = [
    ("NFR-01", "Performance", "โหลดโมเดลครั้งเดียวตอนเริ่ม; ต่อภาพบน GPU: Stage 1 ~208 ms, Stage 2 inference ~10 ms/region + NMS. ครั้งแรกบน UI อาจใช้ ~10–20 วิ (โหลดโมเดล)"),
    ("NFR-02", "Hardware / Platform", "Python 3.11; NVIDIA GPU ≥ 6 GB VRAM (RTX 3050) แนะนำ, CPU รันได้แต่ช้า; Windows 11 (ฟอนต์ไทย tahoma.ttf)"),
    ("NFR-03", "Accuracy", "Stage 2 (train-gray-s) test mAP50 0.853 / mAP50-95 0.537 / recall 0.809; เน้น recall (พลาดตำหนิ = แย่กว่าเตือนเกิน); crazing/crack ยังอ่อนสุด"),
    ("NFR-04", "Usability", "UI ภาษาไทย, flow เดียว: อัปโหลด → กด \"ตรวจสอบ\" → เห็นผล; มี progress indicator; มีภาพตัวอย่างให้ลอง; เน้นความเสี่ยงสูงในสรุปผล"),
    ("NFR-05", "Portability / Deployment", "รันบนเครื่องเดียว ไม่มี server / ฐานข้อมูล / cloud; ต้องต่อเน็ตเฉพาะครั้งแรก (ดาวน์โหลด yolo11s.pt)"),
    ("NFR-06", "Maintainability", "โค้ด functional แยกโมดูล (pipeline เป็นแกน), ตรรกะ fallback รวมที่เดียว (build_regions), มี test_smoke.py, ทุกสคริปต์เตรียมข้อมูลทำซ้ำได้ (reproducible, seed=0), fix_labels สำรอง labels_raw/"),
    ("NFR-07", "Reliability", "fallback เมื่อ Stage 1 หาโลหะไม่เจอ; guard ภาพเปิดไม่ได้ / crop เล็กเกินไป; per-class threshold ปิดได้ด้วย --no-class-conf"),
    ("NFR-08", "Compatibility (train/serve)", "auto-detect ว่าโมเดล Stage 2 เทรนบน grayscale (จาก ckpt.train_args.data) แล้วแปลง input ให้ตรง — กัน train/serve skew"),
    ("NFR-09", "Security / Privacy", "ประมวลผลบนเครื่อง ไม่ส่งภาพออกนอก; ไม่มีบัญชีผู้ใช้/การยืนยันตัวตน (prototype ใช้ในเครื่องเดียว)"),
    ("NFR-10", "Scope / Limitation", "เป็น prototype ผู้ช่วยคัดกรอง (screening assistant) — ไม่รองรับ real-time สายพาน, ไม่ต่อ PLC, ไม่ใช่ระบบตรวจสอบติดตั้งใช้งานจริง"),
]

DATA_DICT = [
    ("*_result.json", "RESULT", "image, metal_regions, metal_area_ratio, fallback_full_image, regions[]", "ผลลัพธ์ต่อภาพจาก process_image()"),
    ("  › regions[]", "REGION", "region_id, box_xywh[4], detections[]", "แต่ละบริเวณที่ตรวจ (embedded)"),
    ("  › › detections[]", "DETECTION", "class, confidence, bbox_xywh, bbox_xyxy_crop, bbox_xyxy_global, region_id", "แต่ละตำหนิที่พบ (embedded)"),
    ("โค้ด (pipeline.py)", "DEFECT_CLASS", "id 0–7, name, name_th, risk", "DEFECT_CLASSES + DEFECT_INFO (คงที่)"),
    ("<split>/labels/*.txt", "LABEL_BOX", "class_id, xc, yc, w, h (normalized)", "label รูปแบบ YOLO ต่อภาพเทรน"),
    ("real_test/labels.csv", "REAL_TEST_LABEL", "filename, classes (คั่นด้วย ; — image-level)", "ground truth ระดับภาพสำหรับ evaluate_real.py"),
    ("thresholds.json", "—", "per_class{cls: {conf, f1_opt, ...}}, macro_f1_*", "per-class confidence threshold จาก tune_thresholds.py"),
    ("results/stage2_*.json", "—", "overall{mAP50, mAP50-95, P, R}, per_class[]", "ผล evaluate.py --mode stage2"),
]

SA = [
    ("SA-01", "อินเทอร์เฟซ 2 ทาง: Web UI (Gradio, app.py) และ CLI (pipeline.py __main__) — ทั้งคู่เรียก pipeline core ตัวเดียวกัน"),
    ("SA-02", "Pipeline Core (pipeline.py) เป็นตัวประสาน: โหลดโมเดล, เรียก Stage 1/2, ทำ NMS, กรอง threshold, วาดผล/บันทึก"),
    ("SA-03", "Stage 1 — DMS46 (Apple Dense Material Segmentation, pre-trained TorchScript): resize ด้านยาว 512 + ImageNet normalize → label map 46 วัสดุ → เลือก index 22 (\"Metal\") → binary mask"),
    ("SA-04", "Region extraction (mask_to_boxes): morphology + findContours + กรองด้วยพื้นที่/สัดส่วน + รวมกรอบที่อยู่ติดกัน → list ของ (x, y, w, h)"),
    ("SA-05", "Fallback (build_regions): ถ้า metal_ratio < 0.05 หรือไม่พบกรอบ → เพิ่มกรอบ = ทั้งภาพ เพื่อไม่พลาดตำหนิ"),
    ("SA-06", "Stage 2 — YOLO11s เทรนเองบน dataset รวม 8 คลาส (grayscale, run: train-gray-s) รันตรวจตำหนิบนแต่ละ region"),
    ("SA-07", "Grayscale guard (run_stage2): ตรวจจาก checkpoint ว่าโมเดลเทรนบน grayscale หรือไม่ ถ้าใช่แปลง crop เป็น grayscale ก่อน predict (กัน train/serve skew)"),
    ("SA-08", "Cross-region NMS (cross_region_nms): แปลง bbox เป็นพิกัดภาพเต็ม แล้วตัด detection ซ้ำข้าม region (คลาสเดียวกัน, IoU > 0.5)"),
    ("SA-09", "Per-class threshold (load_class_conf): กรอง detection ที่ confidence ต่ำกว่าเกณฑ์รายคลาสจาก thresholds.json (ถ้ามี)"),
    ("SA-10", "Output: วาดกรอบ + ป้ายไทย + ความเสี่ยง → UI (4 output) หรือไฟล์ *_result.jpg / *_result.json (CLI)"),
    ("SA-11", "External: DMS46_v1.pt (Apple), Ultralytics YOLO runtime + yolo11s.pt, ระบบไฟล์ท้องถิ่น — ไม่มี API เครือข่าย"),
    ("SA-12", "Deployment: เครื่องเดียว (Windows 11 + RTX 3050 6GB, Python 3.11 venv) — ไม่มี server / ฐานข้อมูล / cloud"),
]

DB = [
    ("DB-01", "ไม่มี DBMS (RDBMS/NoSQL/ORM) — เป็น prototype เครื่องเดียว ประมวลผลต่อคำขอ ไม่มีสถานะร่วมหลายผู้ใช้"),
    ("DB-02", "RESULT — ไฟล์ *_result.json ต่อภาพ: image (path), metal_regions, metal_area_ratio, fallback_full_image, regions[]"),
    ("DB-03", "REGION — embedded ใน RESULT.regions[]: region_id (คีย์), box_xywh[4], detections[]"),
    ("DB-04", "DETECTION — embedded ใน REGION.detections[]: class (→ DEFECT_CLASS), confidence, bbox_xywh, bbox_xyxy_crop, bbox_xyxy_global, region_id"),
    ("DB-05", "DEFECT_CLASS — คงที่ในโค้ด (DEFECT_CLASSES + DEFECT_INFO): id 0–7, name, name_th, risk"),
    ("DB-06", "DATASET_IMAGE — ไฟล์ภาพเทรน: filename (คีย์), split {train/valid/test}; width/height อ่านตอนโหลด ไม่เก็บ"),
    ("DB-07", "LABEL_BOX — ไฟล์ <split>/labels/*.txt (YOLO): class_id, xc, yc, w, h (normalized) บรรทัดละกล่อง"),
    ("DB-08", "REAL_TEST_LABEL — real_test/labels.csv: filename, classes (คั่นด้วย ; — ระดับภาพ) สำหรับ evaluate_real.py"),
    ("DB-09", "ไฟล์ config/ผลการทดลอง: data.yaml, data_oversampled.yaml, thresholds.json, results/stage2_*.json, results/stage1_dms46_test.json, _index.json"),
    ("DB-10", "ความสัมพันธ์: RESULT 1–N REGION ; REGION 1–N DETECTION ; DETECTION N–1 DEFECT_CLASS ; DATASET_IMAGE 1–N LABEL_BOX ; LABEL_BOX N–1 DEFECT_CLASS"),
]

UX = [
    ("UX-01", "หน้าเดียว (Gradio gr.Blocks) ที่ http://127.0.0.1:7860 — ไม่มีล็อกอิน ไม่มีหลายหน้า"),
    ("UX-02", "Flow เดียว: อัปโหลดภาพ → (ปรับ confidence / เปิดตรวจละเอียด ถ้าต้อง) → กด \"ตรวจสอบ\" → อ่านผล"),
    ("UX-03", "Layout: แถวบน = 2 คอลัมน์ (ซ้าย input, ขวา output ข้อความ+ตาราง) ; แถวล่าง = ภาพผลลัพธ์ + ภาพ Stage 1"),
    ("UX-04", "โหลดโมเดลครั้งเดียวตอนเปิดแอป (_ensure_models) — ไม่ค้างเงียบ ๆ ตอนกดปุ่มครั้งแรก ; มี gr.Progress ระหว่างประมวลผล"),
    ("UX-05", "สรุปผล (verdict) เน้น \"ความเสี่ยง\" ก่อนชนิด — ไฮไลต์ตำหนิเสี่ยงสูง (rust, crack)"),
    ("UX-06", "ตารางตำหนิเรียงตามความเสี่ยง (สูง→ต่ำ) แล้วตามความมั่นใจ — ผู้ใช้เห็นสิ่งสำคัญก่อน"),
    ("UX-07", "แสดงภาพ Stage 1 (overlay เขียว = เหล็ก, ส้ม = fallback) ให้ผู้ใช้เข้าใจว่าระบบ \"มองเห็นเหล็กตรงไหน\""),
    ("UX-08", "ข้อความ/ป้ายทั้งหมดเป็นภาษาไทย (วาดบนภาพผ่าน PIL + tahoma.ttf เพราะ cv2.putText ไม่รองรับไทย)"),
    ("UX-09", "มี gr.Examples (test_images/) ให้กดลองได้ทันทีโดยไม่ต้องหาไฟล์"),
    ("UX-10", "รองรับภาพหลายรูปแบบ (grayscale/RGBA) — แปลงให้ภายในไม่ให้ผู้ใช้ต้องจัดการเอง"),
    ("UX-11", "ข้อความกำกับชัดเจนว่าเป็น prototype \"ผู้ช่วยคัดกรอง\" ไม่ใช่ระบบตัดสินขั้นสุดท้าย"),
]

UC = [
    ("UC-01", "ผู้ใช้งาน — อัปโหลดภาพเหล็กเข้าระบบ"),
    ("UC-02", "ผู้ใช้งาน — ปรับ confidence threshold / เปิดโหมดตรวจละเอียด"),
    ("UC-03", "ผู้ใช้งาน — สั่งตรวจภาพ (ระบบรัน Stage 1 + Stage 2)  «include» UC-01"),
    ("UC-04", "ผู้ใช้งาน — ดูผล: กรอบ + ชนิดตำหนิ + ระดับความเสี่ยง + ภาพ Stage 1  «include» UC-03"),
    ("UC-05", "ผู้ใช้งาน — บันทึก/ส่งออกผล (ภาพ + JSON)"),
    ("UC-06", "ผู้พัฒนา — เตรียม dataset (merge → resplit → fix_labels → grayscale → oversample)"),
    ("UC-07", "ผู้พัฒนา — เทรนโมเดล Stage 2 (train.py)"),
    ("UC-08", "ผู้พัฒนา — วัดผล mAP / F1 (evaluate.py)"),
    ("UC-09", "ผู้พัฒนา — หา per-class confidence threshold (tune_thresholds.py)"),
    ("UC-10", "ผู้พัฒนา — วัด Stage 1 เชิงตัวเลข / วัดกับภาพถ่ายจริง (evaluate_stage1.py, evaluate_real.py)"),
    ("UC-11", "ผู้พัฒนา — สร้างรูป/เอกสารประกอบ (make_figures / make_diagrams_doc / make_uml_doc / make_design_doc)"),
]

FLOW = [
    ("F-01", "อ่านภาพ (cv2.imread) ; ถ้าเปิดไม่ได้ → จบ"),
    ("F-02", "Stage 1: run_stage1() → metal mask"),
    ("F-03", "build_regions(mask, image.shape, min_metal_ratio) → boxes, meta"),
    ("F-04", "ตัดสินใจ: metal_ratio < 0.05 หรือไม่พบกรอบ? → ใช่: boxes += (0,0,W,H) (fallback) / ไม่ใช่: ใช้ boxes เดิม"),
    ("F-05", "Pass 1 — วนแต่ละ region: crop ภาพ"),
    ("F-06", "run_stage2(): [ถ้าโมเดล grayscale → แปลง crop เป็น gray] · YOLO predict · กรอง score < class_conf[cls]"),
    ("F-07", "แปลง bbox_xyxy_crop → bbox_xyxy_global (บวก offset ของ region)"),
    ("F-08", "รวม detection ทุก region → cross_region_nms(flat, iou=0.5)"),
    ("F-09", "Pass 2 — วนแต่ละ region: ตัดสินใจ มี detection ที่รอด NMS?"),
    ("F-10", "   ใช่ → วาดกรอบเขียว + กรอบแดง + ป้ายไทย + เก็บลง rows[] / ไม่ → วาดกรอบเขียว + ป้าย \"เหล็ก #k ปกติ\""),
    ("F-11", "CLI: เขียน *_result.jpg + *_result.json (โหมด --folder เพิ่ม _index.json) — UI: คืน (annotated, stage1_img, verdict, rows)"),
]


def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Tahoma"
    doc.styles["Normal"].font.size = Pt(11)

    def h1(t):
        doc.add_page_break()
        doc.add_heading(t, level=1)

    def img(name, width=6.4, caption=None):
        p = DIA / name if (DIA / name).exists() else BASE / "figures" / name
        if not p.exists():
            doc.add_paragraph(f"[ไม่พบรูป {name} — รัน make_diagrams_doc.py / make_uml_doc.py ก่อน]")
            return
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Inches(width))
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size = Pt(9)
            c.runs[0].font.italic = True

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except KeyError:
            t.style = "Table Grid"
        for i, hh in enumerate(headers):
            t.rows[0].cells[i].text = hh
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        return t

    def items(pairs):
        """แต่ละข้อ: (รหัส, ข้อความ) -> bullet ที่ขึ้นต้นด้วยรหัสตัวหนา"""
        for code, txt in pairs:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(code + "  ")
            r.bold = True
            p.add_run(txt)

    # ---- ปก ----
    ti = doc.add_heading("เอกสารออกแบบระบบ (System Design Document)", 0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage: DMS46 → YOLO11s)"
                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("สร้างจาก make_design_doc.py — รูปสถาปัตยกรรม/UML ดึงจาก make_diagrams_doc.py "
                      "และ make_uml_doc.py").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(
        "ขอบเขต: prototype ผู้ช่วยคัดกรองสภาพผิวเหล็กก่อนรับเข้าคลัง/ก่อนตัดใช้งาน "
        "(สเกลโกดัง / ร้านเหล็ก / โรงกลึง) — พนักงานถ่ายรูปด้วยมือถือ ระบบคัดกรองว่ามีตำหนิชนิดใด. "
        "ไม่ใช่ระบบตรวจสอบติดตั้งใช้งานจริงบนสายพานผลิต.")

    # ---- 1. Functional Requirements ----
    h1("1. Functional Requirements")
    doc.add_paragraph("ความสามารถที่ระบบต้องทำได้ (จำแนกตามผู้ใช้: ผู้ใช้งาน = FR-01…FR-10, ผู้พัฒนา/ดูแลโมเดล = FR-11…FR-14)")
    table(["รหัส", "ความต้องการ", "รายละเอียด", "ความสำคัญ"], FR,
          widths=[0.6, 1.6, 4.0, 0.8])

    # ---- 2. Non-functional Requirements ----
    h1("2. Non-functional Requirements")
    table(["รหัส", "ด้าน", "เกณฑ์ / รายละเอียด"], NFR, widths=[0.7, 1.5, 4.8])

    # ---- 3. System Architecture ----
    h1("3. System Architecture")
    doc.add_paragraph("องค์ประกอบสถาปัตยกรรม (แยกเป็นข้อ):")
    items(SA)
    doc.add_heading("3.1 ภาพรวม 2-Stage", level=2)
    img("d1_architecture.png", 5.6)
    doc.add_heading("3.2 Stage 1 — DMS46 Metal Localization (ภายใน)", level=2)
    img("d2_stage1.png", 5.4)
    doc.add_heading("3.3 Stage 2 — YOLO11 Defect Detection (ภายใน)", level=2)
    img("d3_stage2.png", 5.6)
    doc.add_heading("3.4 C4 — System Context (Level 1)", level=2)
    img("uml_3_c4_l1.png", 6.4)
    doc.add_heading("3.5 C4 — Container (Level 2)", level=2)
    img("uml_4_c4_l2.png", 6.4)
    doc.add_heading("3.6 Deployment", level=2)
    img("uml_10_deployment.png", 6.2)

    # ---- 4. Database / Data Design ----
    h1("4. Database Design (Data Design)")
    doc.add_paragraph("รายการข้อมูลที่ระบบเก็บ (แยกเป็นข้อ — ไม่มี DBMS, persist เป็นไฟล์ทั้งหมด):")
    items(DB)
    doc.add_heading("4.1 Conceptual Data Model (ERD)", level=2)
    img("uml_5_erd.png", 6.5,
        "RESULT 1–N REGION 1–N DETECTION N–1 DEFECT_CLASS ; DATASET_IMAGE 1–N LABEL_BOX N–1 DEFECT_CLASS")
    doc.add_heading("4.2 Data Dictionary — ไฟล์และโครงสร้าง", level=2)
    table(["ไฟล์ / แหล่ง", "เอนทิตี", "ฟิลด์หลัก", "หมายเหตุ"], DATA_DICT,
          widths=[1.5, 1.1, 3.0, 1.4])

    # ---- 5. UI/UX Design ----
    h1("5. UI/UX Design")
    doc.add_paragraph("ข้อกำหนดและการออกแบบ UI (แยกเป็นข้อ):")
    items(UX)
    doc.add_heading("5.1 Wireframe", level=2)
    img("ui_wireframe.png", 6.5)
    doc.add_heading("5.2 องค์ประกอบ UI ทีละส่วน", level=2)
    table(["ส่วน", "Component", "หน้าที่ / พฤติกรรม"], [
        ("หัวเรื่อง", "gr.Markdown", "ชื่อระบบ + อธิบาย 2-stage + คลาส 8 ประเภท + ข้อความว่าเป็น prototype"),
        ("ภาพนำเข้า", "gr.Image (numpy)", "อัปโหลด/วาง/กล้อง; รองรับ grayscale/RGBA (แปลงเป็น BGR ภายใน)"),
        ("Confidence", "gr.Slider 0.1–0.9", "threshold รวมของ Stage 2 (ค่าเริ่ม 0.4); เป็นค่า \"ขั้นต่ำ\" เมื่อใช้ per-class threshold"),
        ("ตรวจละเอียด", "gr.Checkbox", "เปิด test-time augmentation (ช้าลง ~2–3x)"),
        ("ปุ่มตรวจสอบ", "gr.Button (primary)", "เรียก analyze(image, conf, detailed); มี gr.Progress แสดงสถานะ"),
        ("ตัวอย่างภาพ", "gr.Examples", "รูปจาก test_images/ ให้กดลองได้ทันที"),
        ("สรุปผล", "gr.Markdown (verdict)", "\"พบตำหนิ N ชนิด\" / \"ความเสี่ยงสูง: ...\" / \"ไม่พบตำหนิ\" + Stage1 % + จำนวนจุด + device"),
        ("ตารางตำหนิ", "gr.Dataframe", "บริเวณ | class | ชนิด(ไทย) | ความมั่นใจ | ความเสี่ยง — เรียงตามความเสี่ยง (สูง→ต่ำ) แล้วความมั่นใจ"),
        ("ภาพผลลัพธ์", "gr.Image", "ภาพต้นฉบับ + กรอบเขียว (region) + กรอบแดง (ตำหนิ) + ป้ายไทย"),
        ("ภาพ Stage 1", "gr.Image", "overlay สีเขียว = พื้นที่เหล็ก, กรอบส้ม = fallback ตรวจทั้งภาพ"),
    ], widths=[1.1, 1.5, 4.0])

    # ---- 6. Flowchart / Use Case ----
    h1("6. Flowchart / Use Case")
    doc.add_heading("6.1 Use Case — แยกเป็นข้อ", level=2)
    items(UC)
    img("uml_1_usecase.png", 6.4, "Use Case Diagram")
    doc.add_heading("6.2 Flowchart — การประมวลผล 1 ภาพ (process_image / analyze)", level=2)
    doc.add_paragraph("ลำดับขั้น (แยกเป็นข้อ):")
    items(FLOW)
    img("uml_8_activity.png", 4.7, "Activity Diagram")
    doc.add_heading("6.3 Flowchart — การเตรียมข้อมูลและเทรน", level=2)
    img("d4_data_prep.png", 6.4)
    doc.add_heading("6.4 State — วงจรชีวิตของภาพระหว่างประมวลผล", level=2)
    img("uml_9_state.png", 6.4)

    # ---- ภาคผนวก ----
    h1("ภาคผนวก — เอกสารประกอบ")
    doc.add_paragraph("• docs/system_diagrams.docx — ไดอาแกรมสถาปัตยกรรม 5 รูป (รายละเอียดภายในแต่ละ Stage + eval)")
    doc.add_paragraph("• docs/uml_sa_diagrams.docx — UML/SA 10 รูป (use case, context, C4 L1/L2, ERD, class, sequence, activity, state, deployment)")
    doc.add_paragraph("• README.md / prepare_data.md / NEXT_STEPS.md / thesis_notes.md — คู่มือและบันทึกการทดลอง")
    doc.add_paragraph("• figures/ — training_curves, per_class_map, class_distribution, confusion_compare")

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX))
    print("\nเขียน:", DOCX.relative_to(BASE))


if __name__ == "__main__":
    print("สร้าง wireframe + เอกสารออกแบบ...")
    make_wireframe()
    build_docx()
