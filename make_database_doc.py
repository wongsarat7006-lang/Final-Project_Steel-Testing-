"""
สร้างเอกสารออกแบบฐานข้อมูล (แยกจาก design_document.docx) -> docs/database_design.docx

    python make_database_doc.py

หัวข้อ: ภาพรวม/เหตุผลการออกแบบ, Conceptual Data Model (ERD), รายการเอนทิตี,
        Data Dictionary (ระดับฟิลด์), ตัวอย่างข้อมูลจริง, Relational Schema เทียบเท่า
เนื้อหาอิงโค้ดจริง (pipeline.py: process_image/DEFECT_CLASSES/DEFECT_INFO) +
ตัวอย่างจริงจาก pipeline_results/ — รูป ERD ดึงจาก make_uml_doc.py (uml_5_erd.png)
ต้องมี: python-docx
"""
import json
from pathlib import Path

BASE = Path(__file__).resolve().parent
DIA = BASE / "figures" / "diagrams"
DOCX = BASE / "docs" / "database_design.docx"

# ตัวอย่างข้อมูลจริง — ผลลัพธ์รันจริงของ pipeline.py บน test_images/rust_example.jpg
SAMPLE_RESULT_JSON = {
    "image": "test_images/rust_example.jpg",
    "metal_regions": 1,
    "metal_area_ratio": 0.4466,
    "fallback_full_image": False,
    "regions": [
        {
            "region_id": 0,
            "box_xywh": [70, 12, 324, 404],
            "detections": [
                {
                    "class": "rust",
                    "confidence": 0.9278,
                    "bbox_xywh": [143.7, 176.0, 271.3, 351.2],
                    "bbox_xyxy_crop": [8.0, 0.4, 279.4, 351.6],
                    "region_id": 0,
                    "bbox_xyxy_global": [78.0, 12.4, 349.4, 363.6],
                }
            ],
        }
    ],
}

# ---------- 1. ภาพรวม ----------
OVERVIEW = [
    "ระบบเป็น prototype รันเครื่องเดียว (single-machine) ประมวลผลทีละคำขอ ไม่มีผู้ใช้พร้อมกันหลายคน "
    "และไม่มีความต้องการ transaction/ACID ข้ามคำขอ — ตาม NFR-05 (Portability) และ NFR-09 (Security/Privacy: "
    "ประมวลผลบนเครื่อง ไม่ส่งข้อมูลออกนอก) จึง **ไม่ติดตั้ง DBMS** (ไม่มี RDBMS/NoSQL/ORM ใดๆ) และไม่มี server",
    "ข้อมูลทั้งหมด persist เป็น **ไฟล์บนดิสก์** แทนตาราง: ผลตรวจต่อภาพเป็น JSON (`*_result.json`), "
    "ป้ายกำกับชุดเทรนเป็น YOLO .txt, ป้ายกำกับชุดทดสอบจริงเป็น .csv, ค่าคงที่ (นิยามคลาส) ฝังในโค้ด",
    "เหตุผลที่เลือก JSON แบบ nested (ฝัง REGION/DETECTION ไว้ใน RESULT ไฟล์เดียว) แทนตารางแยก: "
    "อ่าน-เขียนต่อภาพเป็นหน่วยเดียวพอดี (1 คำขอ = 1 ไฟล์ผลลัพธ์), ไม่ต้อง join, ไม่มีการอัปเดตย้อนหลัง — "
    "เป็นข้อมูลที่เขียนครั้งเดียวจบ (write-once, append-only ต่อการรัน)",
    "ถ้าอนาคตต้องขยายเป็นระบบหลายผู้ใช้ / เก็บประวัติการตรวจระยะยาว / query ข้ามภาพ (เช่น \"ดูภาพที่พบ rust ทั้งหมดในเดือนนี้\") "
    "จึงจะคุ้มที่จะย้ายเข้า RDBMS จริง — ดูข้อ 6 (Relational Schema เทียบเท่า) สำหรับแนวทางถ้าต้องย้าย",
]

# ---------- 2. รายการเอนทิตี (แยกเป็นข้อ) ----------
DB = [
    ("DB-01", "ไม่มี DBMS (RDBMS/NoSQL/ORM) — เป็น prototype เครื่องเดียว ประมวลผลต่อคำขอ ไม่มีสถานะร่วมหลายผู้ใช้"),
    ("DB-02", "RESULT — ไฟล์ *_result.json ต่อภาพ: image (path), metal_regions, metal_area_ratio, fallback_full_image, regions[]"),
    ("DB-03", "REGION — embedded ใน RESULT.regions[]: region_id (คีย์), box_xywh[4], detections[]"),
    ("DB-04", "DETECTION — embedded ใน REGION.detections[]: class (→ DEFECT_CLASS), confidence, bbox_xywh, bbox_xyxy_crop, bbox_xyxy_global, region_id"),
    ("DB-05", "DEFECT_CLASS — คงที่ในโค้ด (DEFECT_CLASSES + DEFECT_INFO): id 0–7, name, name_th, risk"),
    ("DB-06", "DATASET_IMAGE — ไฟล์ภาพเทรน: filename (คีย์), split {train/valid/test}; width/height อ่านตอนโหลด ไม่เก็บ"),
    ("DB-07", "LABEL_BOX — ไฟล์ <split>/labels/*.txt (YOLO): class_id, xc, yc, w, h (normalized) บรรทัดละกล่อง"),
    ("DB-08", "REAL_TEST_LABEL — real_test/labels.csv: filename, classes (คั่นด้วย ; — ระดับภาพ) สำหรับ evaluate_real.py"),
    ("DB-09", "ไฟล์ config/ผลการทดลอง: data.yaml, data_oversampled.yaml, thresholds.json, results/stage2_*.json, results/stage1_dms46_test.json, _index.json"),
    ("DB-10", "ความสัมพันธ์: RESULT 1–N REGION ; REGION 1–N DETECTION ; DETECTION N–1 DEFECT_CLASS ; DATASET_IMAGE 1–N LABEL_BOX ; LABEL_BOX N–1 DEFECT_CLASS"),
]

# ---------- 3. Data Dictionary ระดับฟิลด์ ----------
FIELDS = [
    ("RESULT", "image", "string (path)", "path ของภาพต้นฉบับที่ตรวจ"),
    ("RESULT", "metal_regions", "int", "จำนวนบริเวณที่ Stage 1 พบ (รวมกรอบ fallback ถ้ามี)"),
    ("RESULT", "metal_area_ratio", "float 0–1", "สัดส่วนพื้นที่ที่ Stage 1 ระบุว่าเป็นโลหะ ต่อพื้นที่ภาพทั้งหมด"),
    ("RESULT", "fallback_full_image", "bool", "true ถ้า metal_area_ratio < 0.05 หรือไม่พบบริเวณเลย (เพิ่มกรอบทั้งภาพ)"),
    ("RESULT", "regions[]", "array<REGION>", "รายการบริเวณที่ตรวจ (embedded)"),
    ("REGION", "region_id", "int", "ลำดับบริเวณในภาพนี้ (0-based)"),
    ("REGION", "box_xywh[4]", "int×4", "กรอบบริเวณในพิกัดภาพเต็ม (x, y, w, h)"),
    ("REGION", "detections[]", "array<DETECTION>", "ตำหนิที่พบในบริเวณนี้ (embedded, ว่างได้ = ปกติ)"),
    ("DETECTION", "class", "string (FK→DEFECT_CLASS.name)", "ชื่อคลาสภาษาอังกฤษ 1 ใน 8 ประเภท"),
    ("DETECTION", "confidence", "float 0–1", "ความมั่นใจจาก YOLO หลังกรอง per-class threshold"),
    ("DETECTION", "bbox_xywh[4]", "float×4", "กรอบตำหนิในพิกัดของ crop (center-x, center-y, w, h)"),
    ("DETECTION", "bbox_xyxy_crop[4]", "float×4", "กรอบตำหนิในพิกัดของ crop (x1, y1, x2, y2)"),
    ("DETECTION", "bbox_xyxy_global[4]", "float×4", "กรอบตำหนิแปลงกลับเป็นพิกัดภาพเต็ม (= crop + offset ของ region)"),
    ("DETECTION", "region_id", "int (FK→REGION.region_id)", "อ้างอิงกลับไปยังบริเวณต้นทาง (ไว้ใช้ตอน cross-region NMS)"),
    ("DEFECT_CLASS", "id", "int 0–7 (PK)", "ลำดับ index ตามที่โมเดล YOLO เอาต์พุต"),
    ("DEFECT_CLASS", "name", "string", "ชื่อคลาสภาษาอังกฤษ (ตรงกับ label ชุดเทรน)"),
    ("DEFECT_CLASS", "name_th", "string", "ชื่อไทยไว้แสดงผลบน UI/ภาพ"),
    ("DEFECT_CLASS", "risk", "string (enum)", "ระดับความเสี่ยง: ต่ำ / ต่ำ-ปานกลาง / ปานกลาง / ปานกลาง-สูง / สูง"),
    ("DATASET_IMAGE", "filename", "string (PK)", "ชื่อไฟล์ภาพในชุดเทรน/วัดผล"),
    ("DATASET_IMAGE", "split", "string (enum)", "train / valid / test"),
    ("LABEL_BOX", "class_id", "int (FK→DEFECT_CLASS.id)", "คลาสของกล่อง"),
    ("LABEL_BOX", "xc, yc, w, h", "float×4 (0–1)", "พิกัด normalized ตามรูปแบบ YOLO"),
    ("REAL_TEST_LABEL", "filename", "string", "ชื่อไฟล์ภาพใน real_test/images/"),
    ("REAL_TEST_LABEL", "classes", "string (multi-value, คั่น ;)", "ชนิดตำหนิที่เห็นจริงในภาพ ระดับภาพ (ไม่ตีกรอบ) — ดูหมายเหตุ normalization ข้อ 6"),
]

# ---------- 6. Relational Schema เทียบเท่า (ถ้าย้ายไป RDBMS) ----------
RELATIONAL = [
    ("RESULT", "result_id PK, image_path, metal_regions, metal_area_ratio, fallback_full_image, created_at"),
    ("REGION", "region_id PK, result_id FK→RESULT, box_x, box_y, box_w, box_h"),
    ("DETECTION", "detection_id PK, region_id FK→REGION, class_id FK→DEFECT_CLASS, confidence, "
                  "bbox_x1, bbox_y1, bbox_x2, bbox_y2 (global)"),
    ("DEFECT_CLASS", "class_id PK, name, name_th, risk"),
    ("DATASET_IMAGE", "image_id PK, filename UNIQUE, split"),
    ("LABEL_BOX", "label_id PK, image_id FK→DATASET_IMAGE, class_id FK→DEFECT_CLASS, xc, yc, w, h"),
    ("REAL_TEST_IMAGE", "image_id PK, filename UNIQUE"),
    ("REAL_TEST_LABEL_CLASS", "image_id FK→REAL_TEST_IMAGE, class_id FK→DEFECT_CLASS  (PK คู่ — junction table)"),
]


def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Tahoma"
    doc.styles["Normal"].font.size = Pt(11)

    def h1(t):
        doc.add_page_break()
        doc.add_heading(t, level=1)

    def img(name, width=6.4, caption=None):
        p = DIA / name if (DIA / name).exists() else BASE / "figures" / name
        if not p.exists():
            doc.add_paragraph(f"[ไม่พบรูป {name} — รัน make_uml_doc.py ก่อน]")
            return
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Inches(width))
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size = Pt(9)
            c.runs[0].font.italic = True

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except KeyError:
            t.style = "Table Grid"
        for i, hh in enumerate(headers):
            t.rows[0].cells[i].text = hh
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        return t

    def items(pairs):
        """แต่ละข้อ: (รหัส, ข้อความ) -> bullet ที่ขึ้นต้นด้วยรหัสตัวหนา"""
        for code, txt in pairs:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(code + "  ")
            r.bold = True
            p.add_run(txt)

    def code_block(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        return p

    # ---- ปก ----
    ti = doc.add_heading("เอกสารออกแบบฐานข้อมูล (Database Design Document)", 0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ระบบตรวจจับตำหนิพื้นผิวเหล็ก (2-Stage: DMS46 → YOLO11s)"
                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("สร้างจาก make_database_doc.py — แยกจาก design_document.docx หัวข้อ 4 เดิม "
                      "พร้อมรายละเอียดระดับฟิลด์ + ตัวอย่างข้อมูลจริง + แนวทาง schema เชิงสัมพันธ์"
                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ---- 1. ภาพรวม ----
    h1("1. ภาพรวมและเหตุผลการออกแบบ")
    for p in OVERVIEW:
        doc.add_paragraph(p, style="List Bullet")

    # ---- 2. Conceptual Data Model (ERD) ----
    h1("2. Conceptual Data Model (ERD)")
    doc.add_paragraph("ความสัมพันธ์ระหว่างเอนทิตี (แยกเป็นข้อ):")
    items(DB)
    doc.add_heading("2.1 แผนภาพ ERD", level=2)
    img("uml_5_erd.png", 6.5,
        "RESULT 1–N REGION 1–N DETECTION N–1 DEFECT_CLASS ; DATASET_IMAGE 1–N LABEL_BOX N–1 DEFECT_CLASS")

    # ---- 3. Data Dictionary ----
    h1("3. Data Dictionary (ระดับฟิลด์)")
    doc.add_paragraph("รายละเอียดทุกฟิลด์ของแต่ละเอนทิตี อิงจากโครงสร้างจริงใน pipeline.py "
                      "(process_image, DEFECT_CLASSES, DEFECT_INFO) และไฟล์ label จริง:")
    table(["เอนทิตี", "ฟิลด์", "ชนิดข้อมูล", "คำอธิบาย"], FIELDS,
          widths=[1.1, 1.3, 1.6, 2.6])

    # ---- 4. ไฟล์และแหล่งข้อมูลอื่น ----
    h1("4. ไฟล์ข้อมูลอื่นในระบบ")
    table(["ไฟล์ / แหล่ง", "เอนทิตี", "ฟิลด์หลัก", "หมายเหตุ"], [
        ("*_result.json", "RESULT", "image, metal_regions, metal_area_ratio, fallback_full_image, regions[]", "ผลลัพธ์ต่อภาพจาก process_image()"),
        ("<split>/labels/*.txt", "LABEL_BOX", "class_id, xc, yc, w, h (normalized)", "label รูปแบบ YOLO ต่อภาพเทรน"),
        ("real_test/labels.csv", "REAL_TEST_LABEL", "filename, classes (คั่นด้วย ;)", "ground truth ระดับภาพสำหรับ evaluate_real.py"),
        ("thresholds.json", "—", "per_class{cls: {conf, f1_opt, ...}}, macro_f1_*", "per-class confidence threshold จาก tune_thresholds.py"),
        ("results/stage2_*.json", "—", "overall{mAP50, mAP50-95, P, R}, per_class[]", "ผล evaluate.py --mode stage2"),
        ("data.yaml / data_oversampled.yaml", "—", "path, train/val/test, nc, names", "config ชุดข้อมูลของ Ultralytics YOLO"),
    ], widths=[1.7, 1.1, 2.7, 1.6])

    # ---- 5. ตัวอย่างข้อมูลจริง ----
    h1("5. ตัวอย่างข้อมูลจริง (Sample Data)")
    doc.add_paragraph(
        "ผลลัพธ์จริงจากการรัน pipeline.py บนภาพ test_images/rust_example.jpg "
        "(1 บริเวณเหล็ก พบตำหนิ 1 จุด — rust ความมั่นใจ 92.8%):"
    )
    code_block(json.dumps(SAMPLE_RESULT_JSON, ensure_ascii=False, indent=2))

    # ---- 6. Relational Schema เทียบเท่า ----
    h1("6. Relational Schema เทียบเท่า (แนวทางถ้าย้ายไปใช้ RDBMS ในอนาคต)")
    doc.add_paragraph(
        "ระบบปัจจุบันไม่ใช้ DBMS (ข้อ 1) แต่ถ้าต้องขยายเป็นระบบเก็บประวัติ/หลายผู้ใช้ในอนาคต "
        "ตารางด้านล่างคือแนวทางแปลงเอนทิตีปัจจุบันเป็นสคีมาเชิงสัมพันธ์ (3NF):"
    )
    table(["ตาราง", "คอลัมน์ (PK/FK ระบุไว้)"], RELATIONAL, widths=[1.7, 4.7])
    doc.add_heading("6.1 หมายเหตุการทำ Normalization", level=2)
    doc.add_paragraph(
        "REAL_TEST_LABEL ปัจจุบันเก็บ `classes` เป็นสตริงหลายค่าคั่นด้วย ';' ในไฟล์เดียว (ผิด 1NF) "
        "เพราะออกแบบมาให้ผู้ใช้แก้ด้วย Excel/Notepad ได้ง่ายที่สุด — ถ้าย้ายเข้า RDBMS จริง ควรแยกเป็น "
        "REAL_TEST_IMAGE (1 แถวต่อภาพ) และ REAL_TEST_LABEL_CLASS (junction table ความสัมพันธ์ N–N "
        "ระหว่างภาพกับคลาสตำหนิ) ตามตารางข้างต้น เพื่อ query ระดับคลาสได้ตรงไปตรงมา "
        "(เช่น \"นับภาพที่มี rust ทั้งหมด\" โดยไม่ต้อง parse string)."
    )

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX))
    print("เขียน:", DOCX.relative_to(BASE))


if __name__ == "__main__":
    build_docx()
