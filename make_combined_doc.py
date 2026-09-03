"""
รวม docs/design_document.docx + docs/system_diagrams.docx + docs/uml_sa_diagrams.docx
เป็นไฟล์เดียว -> docs/combined_document.docx (ปกรวม + สารบัญ + ทั้ง 3 ส่วน คั่นด้วยหน้าใหม่)

    python make_combined_doc.py

ต้องมี: python-docx, docxcompose (pip install docxcompose)
รัน make_design_doc.py / make_diagrams_doc.py / make_uml_doc.py ก่อน ถ้ายังไม่มี docs/*.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

BASE = Path(__file__).resolve().parent
DOCS = BASE / "docs"
OUT = DOCS / "combined_document.docx"

PARTS = [
    ("ส่วนที่ 1 — เอกสารออกแบบระบบ", "Part 1 — System Design Document",
     DOCS / "design_document.docx"),
    ("ส่วนที่ 2 — ไดอาแกรมสถาปัตยกรรมระบบ", "Part 2 — System Architecture Diagrams",
     DOCS / "system_diagrams.docx"),
    ("ส่วนที่ 3 — UML / SA Diagrams", "Part 3 — UML / SA Diagrams",
     DOCS / "uml_sa_diagrams.docx"),
]

ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_toc_field(doc):
    """แทรก TOC field — เปิดใน Word แล้วกด F9 (หรือคลิกขวา > Update Field) เพื่อสร้างสารบัญจริง"""
    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)

    hint_run = p.add_run("คลิกขวาที่นี่ > Update Field เพื่อสร้างสารบัญ (หรือกด F9)")
    hint_run.italic = True
    hint_run.font.color.rgb = GREY

    end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def add_page_numbers(section):
    """ใส่เลขหน้าไว้กลางท้ายกระดาษ ให้อ่านง่ายเวลาพิมพ์/อ้างอิง"""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def main():
    for label, label_en, path in PARTS:
        if not path.exists():
            raise SystemExit(f"ไม่พบ {path} — รันสคริปต์ที่สร้างไฟล์นั้นก่อน")

    master = Document()
    master.styles["Normal"].font.name = "Tahoma"
    master.styles["Normal"].font.size = Pt(11)

    section = master.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)
    add_page_numbers(section)

    # ---------- ปกรวม ----------
    for _ in range(3):
        master.add_paragraph()
    title = master.add_heading(
        "ระบบตรวจจับตำหนิพื้นผิวเหล็กด้วยภาพถ่าย (2-Stage)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = master.add_paragraph("Steel Surface Defect Detection System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(15)
    sub.runs[0].font.color.rgb = GREY
    sub.runs[0].italic = True

    tag = master.add_paragraph("เอกสารประกอบโครงการฉบับรวม  —  Combined Project Documentation")
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.runs[0].bold = True
    tag.runs[0].font.size = Pt(12)
    tag.runs[0].font.color.rgb = ACCENT

    for _ in range(3):
        master.add_paragraph()

    lead = master.add_paragraph("เอกสารฉบับนี้รวมเนื้อหาทั้งหมด 3 ส่วน:")
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.runs[0].bold = True

    for label, label_en, _ in PARTS:
        li = master.add_paragraph()
        li.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = li.add_run(f"{label}")
        r1.font.size = Pt(12)
        r2 = li.add_run(f"   ({label_en})")
        r2.font.size = Pt(10)
        r2.italic = True
        r2.font.color.rgb = GREY

    master.add_page_break()

    # ---------- สารบัญ ----------
    master.add_heading("สารบัญ (Table of Contents)", level=1)
    add_toc_field(master)
    master.add_page_break()

    # ---------- รวม 3 ส่วน ----------
    composer = Composer(master)
    for i, (label, label_en, path) in enumerate(PARTS):
        if i > 0:
            master.add_page_break()
        p = master.add_heading("", level=1)
        r = p.add_run(label)
        r.font.color.rgb = ACCENT
        r2 = p.add_run(f"   ({label_en})")
        r2.italic = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = GREY

        sub_doc = Document(str(path))
        composer.append(sub_doc)

    DOCS.mkdir(parents=True, exist_ok=True)
    composer.save(str(OUT))
    print(f"บันทึกแล้ว: {OUT}")


if __name__ == "__main__":
    main()
